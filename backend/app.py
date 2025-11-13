# ============================================================================
# EVOLV TRAINING COPILOT API — LABELED EDITION
# ============================================================================
# This file is the same code as your original app.py, with *extra comments*
# that act as "handles" so we can tell you **exactly** where to paste future
# patches. Nothing functional has changed.
#
# HOW TO USE THE LABELS
# ---------------------
# - Big banner sections use bracketed IDs like: [S1], [S2], [S3A], etc.
# - Subsections and utilities use: [Sx.y], [UT-...], [CFG-...], etc.
# - When creating new areas use subsections and utility codes.
# - Every function has a lightweight label: [FN:<name>]
# - When I give you instructions later, I’ll say things like:
#     "Paste under [S7] ROUTES: Q&A Admin — List/Add/Update/Delete"
#   or "Insert this function under [S5.3] Helpers — Facility detection utilities"
#
# QUICK TABLE OF CONTENTS (search these exact tags in this file)
# --------------------------------------------------------------
# [S0]  File Preamble & Imports
# [S1]  Speed Cache Utilities
# [S2]  Environment & Core Config
# [S3]  Abbreviations (Paths, CSV Helpers, Map/Normalization)
# [S4]  Clients & Collections (OpenAI/Chroma)
# [S5]  FastAPI App & CORS
# [S6]  Models, Prompts & Ask Payload
# [S7]  Helpers — Admin Token check, CSV load/save, Embeddings, MMR, etc.
# [S5.3] Helpers — Facility detection utilities
# [S8]  Retrieval Pipeline (multi-query + MMR + boosting)
# [S9]  User Q&A Log (CSV) + Admin routes for log
# [S10] Chroma Safety & Rebuild (QA & Docs)
# [S11] Docs Rebuild from /docs folder
# [S12] Health & Root Routes (simple)
# [S13] Rule-based Answer Extractor
# [S14] Main /ask Endpoint
# [S15] Q&A Admin — List/Add/Update/Delete/Upload/Reindex
# [S16] Bulk Ingest from Document (PDF/DOCX/TXT) to Q&A
# [S17] Intake Submit (Client Info Spreadsheet -> QA rows + index upsert)
# [S18] Main Entrypoint (__main__)
# [S19] Abbreviation Admin Routes
# ============================================================================


# ==============================================================================
# [S0] PREAMBLE & IMPORTS
# ==============================================================================
import os
import csv
import io
from typing import List

import pandas as pd
from fastapi import FastAPI, HTTPException, Request, UploadFile, File, Form
import secrets
from fastapi import Query
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
import chromadb
from openai import OpenAI

import re


import math
from collections import OrderedDict

import time
from functools import lru_cache

from fastapi.responses import JSONResponse, Response, FileResponse, HTMLResponse
import threading  # <-- add this

from pathlib import Path   # <-- needed for FACILITY_HTML path
import logging


# ==============================================================================
# [S1] SPEED CACHE UTILITIES
# ==============================================================================
# ======== Speed caches ========
_ASK_CACHE = {}                 # (question, section_pref) -> (payload, ts)
_ASK_TTL_SEC = 300              # 5 minutes


# [FN:_ask_cache_get] ----------------------------------------------------
def _ask_cache_get(key):
    item = _ASK_CACHE.get(key)
    if not item: return None
    val, ts = item
    if time.time() - ts > _ASK_TTL_SEC:
        _ASK_CACHE.pop(key, None)
        return None
    return val


# [FN:_ask_cache_put] ----------------------------------------------------
def _ask_cache_put(key, value):
    _ASK_CACHE[key] = (value, time.time())

@lru_cache(maxsize=2048)

# [FN:_embed_cached] ----------------------------------------------------
def _embed_cached(text: str):
    return embed_query(text)


# ==============================================================================
# [S2] ENVIRONMENT & CORE CONFIG
# ==============================================================================
# =========================
# Env & Config
# =========================
load_dotenv()
API_KEY = os.getenv("OPENAI_API_KEY")
if not API_KEY:
    raise RuntimeError("OPENAI_API_KEY missing in .env")

ADMIN_TOKEN = os.getenv("ADMIN_TOKEN", "")
CSV_PATH = os.path.join("data", "qa.csv")

# === SQLite target (we'll migrate into this) ===
DB_PATH = os.path.join("data", "evolv.db")

# Toggle: read from SQLite instead of CSV (writes still go to CSV in Step 2)
USE_SQLITE_READ = os.getenv("USE_SQLITE_READ", "0") == "1"

# Toggle: also mirror all writes into SQLite (safe, Step 3)
MIRROR_SQLITE_WRITES = os.getenv("MIRROR_SQLITE_WRITES", "1") == "1"

# (For Step 4) Toggle: write to SQLite instead of CSV (off for Step 3)
USE_SQLITE_WRITE = os.getenv("USE_SQLITE_WRITE", "0") == "1"



# ==============================================================================
# [S3] ABBREVIATIONS — Paths & CSV helpers
# ==============================================================================
# ---- Abbreviation file path (single source of truth) ----
# === Dictionary (unified) ===
DICT_CSV_PATH = os.path.join("data", "dictionary.csv")

def _ensure_dict_csv():
    os.makedirs(os.path.dirname(DICT_CSV_PATH), exist_ok=True)
    if not os.path.exists(DICT_CSV_PATH):
        with open(DICT_CSV_PATH, "w", newline="", encoding="utf-8") as f:
            f.write("key,canonical,kind,notes,match_mode\n")

def _read_dict_rows_csv() -> list[dict]:
    _ensure_dict_csv()
    out = []
    with open(DICT_CSV_PATH, newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            out.append({
                "key": (row.get("key") or "").strip(),
                "canonical": (row.get("canonical") or "").strip(),
                "kind": (row.get("kind") or "abbr").strip(),
                "notes": (row.get("notes") or "").strip(),
                "match_mode": (row.get("match_mode") or "exact").strip().lower() or "exact",
            })
    return out

def dict_rows_from_sqlite() -> list[dict]:
    return db_fetchall("""
        SELECT key, canonical, kind, notes, COALESCE(NULLIF(match_mode,''),'exact') AS match_mode
        FROM dictionary
        ORDER BY LOWER(key)
    """)

def load_dictionary_unified() -> list[dict]:
    """SQLite is the source of truth; if empty, optionally fall back to CSV."""
    rows = dict_rows_from_sqlite()
    if not rows:
        rows = _read_dict_rows_csv()
    return [r for r in rows if r.get("key") and r.get("canonical")]

@lru_cache(maxsize=1_000)
def _dictionary_maps():
    """
    Build fast lookup structures:
      - exact_map: lower(key) -> canonical
      - word_set: tokens that should normalize at word-level (match_mode='word')
      - facility_aliases: lower(key) -> canonical (kind='facility_alias')
    """
    rows = load_dictionary_unified()
    exact_map = {}
    word_set  = set()
    facility_aliases = {}

    for r in rows:
        k = (r["key"] or "").strip().lower()
        c = (r["canonical"] or "").strip()
        mm = (r.get("match_mode") or "exact").strip().lower()
        kind = (r.get("kind") or "abbr").strip().lower()

        if not k or not c: 
            continue
        exact_map[k] = c
        if mm == "word":
            word_set.add(k)
        if kind == "facility_alias":
            facility_aliases[k] = c

    # compile a boundary pattern for exact keys with punctuation escaped
    if exact_map:
        pat = re.compile(r"\b(" + "|".join(map(re.escape, exact_map.keys())) + r")\b", re.IGNORECASE)
    else:
        pat = None

    return {
        "exact_map": exact_map,
        "word_set":  word_set,
        "facility_aliases": facility_aliases,
        "pattern": pat
    }

def normalize_dictionary_query(text: str) -> str:
    """
    Expand known keys to 'key (canonical)' to boost both exact and semantic retrieval.
    Handles both exact (boundary) replacements and word-level tokens.
    """
    if not text:
        return text
    d = _dictionary_maps()
    pat = d["pattern"]
    out = text

    # 1) exact boundary replacements
    if pat:
        def _repl(m):
            raw = m.group(0)
            canon = d["exact_map"].get(raw.lower())
            return f"{raw} ({canon})" if canon else raw
        out = pat.sub(_repl, out)

    # 2) word-level (token) replacements (only if not already expanded)
    tokens = re.split(r"(\W+)", out)  # keep separators
    for i, tk in enumerate(tokens):
        lk = tk.lower()
        if lk in d["word_set"] and d["exact_map"].get(lk):
            if f"({d['exact_map'][lk]})" not in tk:  # avoid double expansion
                tokens[i] = f"{tk} ({d['exact_map'][lk]})"
    out = "".join(tokens)
    return out


def clear_dictionary_caches():
    _dictionary_maps.cache_clear()

# Hybrid search config
EMBED_MODEL = "text-embedding-3-large"   # must match your index scripts
DOCS_COLLECTION_NAME = "docs_main"       # created by build_docs_index.py
QA_WEIGHT   = 1.0
DOCS_WEIGHT = 0.9

# ==============================================================================
# [S4] CLIENTS & COLLECTIONS (OpenAI / Chroma)
# ==============================================================================
# =========================
# Clients & Collections
# =========================
client = OpenAI(api_key=API_KEY)
chroma = chromadb.PersistentClient(path="vectorstore")

# Q&A collection
collection = chroma.get_or_create_collection(name="qa_main")

# Docs collection
docs_collection = chroma.get_or_create_collection(name=DOCS_COLLECTION_NAME)


# ==============================================================================
# [S5] FASTAPI APP & CORS
# ==============================================================================
# =========================
# FastAPI app
# =========================
app = FastAPI(title="Evolv Training Copilot API")

from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response, FileResponse, HTMLResponse

ALLOWED_ORIGINS = [
    "https://www-startevolv-com.filesusr.com",            # live Admin page (you confirmed)
    "https://454efc33-431a-4209-8b28-8f65021905a4.filesusr.com",  # Wix editor page (you just confirmed)
    "https://www.startevolv.com",                         # your custom domain (future-proof)
    "https://editor.wix.com",                             # Wix editor shell
    "https://*.wixsite.com",
    "https://towery-supervoluminously-chuck.ngrok-free.dev",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    # cover any wixsite.com preview subdomain
    allow_origin_regex=r"^https://([a-z0-9-]+\.)?wixsite\.com$",
    # (optional) cover any filesusr.com subdomain in case Wix rotates IDs:
    # allow_origin_regex also accepts a tuple/list; here’s a combined pattern:
    # allow_origin_regex=r"^(https://([a-z0-9-]+\.)?wixsite\.com|https://[a-z0-9-]+\.filesusr\.com)$",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    max_age=86400,
)

@app.options("/{rest_of_path:path}")

# [FN:cors_preflight_always_ok] ----------------------------------------------------
def cors_preflight_always_ok(rest_of_path: str):
    return Response(status_code=204)


# Let Wix (and your own site) embed this app in an <iframe>
@app.middleware("http")
async def set_frame_ancestors(request, call_next):
    resp = await call_next(request)
    # The sites allowed to frame your app (Wix preview/editor + your domain)
    resp.headers["Content-Security-Policy"] = (
        "frame-ancestors https://*.wixsite.com https://editor.wix.com https://www.startevolv.com"
    )
    return resp



# Log to the default uvicorn logger so you can see path checks in the console
logger = logging.getLogger("uvicorn.error")

# Absolute path to your HTML file, derived relative to this app.py
FACILITY_HTML = Path(__file__).parent.parent / "frontend" / "TEST  Facility_Details.html"
# NOTE: there are **two spaces** between "TEST" and "Facility" above — keep them.

@app.get("/facility", response_class=HTMLResponse)
def serve_facility():
    logger.info(f"Serving facility file: {FACILITY_HTML} (exists={FACILITY_HTML.exists()})")
    if not FACILITY_HTML.exists():
        return HTMLResponse(
            content=f"<h1>404</h1><p>File not found: {FACILITY_HTML}</p>",
            status_code=404,
        )
    return FileResponse(FACILITY_HTML)


# --- Serve Admin & Ask HTML from the same frontend folder ---
ADMIN_HTML = Path(__file__).parent.parent / "frontend" / "TEST Admin.html"
ASK_HTML   = Path(__file__).parent.parent / "frontend" / "TEST Ask.html"

@app.get("/admin", response_class=HTMLResponse)
def serve_admin():
    logger.info(f"Serving admin file: {ADMIN_HTML} (exists={ADMIN_HTML.exists()})")
    if not ADMIN_HTML.exists():
        return HTMLResponse(
            content=f"<h1>404</h1><p>File not found: {ADMIN_HTML}</p>",
            status_code=404,
        )
    return FileResponse(ADMIN_HTML)

@app.get("/ask-page", response_class=HTMLResponse)
def serve_ask_page():
    logger.info(f"Serving ask page file: {ASK_HTML} (exists={ASK_HTML.exists()})")
    if not ASK_HTML.exists():
        return HTMLResponse(
            content=f"<h1>404</h1><p>File not found: {ASK_HTML}</p>",
            status_code=404,
        )
    return FileResponse(ASK_HTML)


# ==============================================================================
# [S5.1] SQLITE — connection + schema init
# ==============================================================================
import sqlite3

def _db():
    import sqlite3, os
    from pathlib import Path
    DB_PATH = Path("data/evolv.db")
    DB_PATH.parent.mkdir(parents=True, exist_ok=True)

    # NOTE: check_same_thread=False allows reuse in the same process/thread pool
    conn = sqlite3.connect(DB_PATH, check_same_thread=False, timeout=10.0)

    # If DB is on a cloud-synced drive, WAL can fail. Allow disabling via env.
    use_wal = (os.environ.get("USE_SQLITE_WAL", "1") == "1")

    try:
        if use_wal:
            conn.execute("PRAGMA journal_mode=WAL;")
        else:
            # Fallback journaling for cloud drives
            conn.execute("PRAGMA journal_mode=DELETE;")
            conn.execute("PRAGMA synchronous=NORMAL;")
    except Exception:
        # If setting PRAGMA fails, keep going with defaults
        pass

    # Avoid lock stalls
    conn.execute("PRAGMA busy_timeout=3000;")
    return conn

# ---- Facilities helpers -------------------------------------------------------
def _gen_token(n: int = 24) -> str:
    """Short, URL-safe token for intake links."""
    return secrets.token_urlsafe(n)[:n]

def _ensure_unique_facility_id(conn, base_slug: str) -> str:
    """Guarantee a unique facility_id by appending -2, -3, ... if needed."""
    slug = base_slug or "facility"
    cur = conn.cursor()
    cur.execute("SELECT 1 FROM facilities WHERE facility_id = ? LIMIT 1;", (slug,))
    if not cur.fetchone():
        return slug
    i = 2
    while True:
        cand = f"{slug}-{i}"
        cur.execute("SELECT 1 FROM facilities WHERE facility_id = ? LIMIT 1;", (cand,))
        if not cur.fetchone():
            return cand
        i += 1

def _facility_intake_urls(request: Request, facility_id: str, token: str) -> dict:
    """Return both full and path-only intake URLs."""
    base = str(request.base_url).rstrip("/")
    path = f"/intake/f/{facility_id}?t={token}"
    return {"full": f"{base}{path}", "path": path}
# -------------------------------------------------------------------------------
# ---- Public Intake routes (token = ?t=...) -----------------------------------
from fastapi import Query

@app.get("/intake/f/{facility_id}")
def intake_get_facility(request: Request, facility_id: str, t: str = Query(default="")):
    """
    Read-only intake fetch. Validates the facility's intake token (?t=...).
    Returns facility columns + contacts + partners so the intake page can prefill.
    """
    import sqlite3 as _sqlite3
    conn = _db()
    cur  = conn.cursor()

    # 1) Check token + get facility row
    cur.execute("""
        SELECT facility_id, facility_name, corporate_group, intake_token, intake_status,
               created_at, updated_at, extras,

               -- All intake fields we store on facilities
               legal_name, address_line1, address_line2, city, state, zip, county,
               avg_dcs, short_beds, ltc_beds, outpatient_pt,
               emr, emr_other, pt_emr, orders, orders_other,
               -- optional JSON snapshots
               raw_json
        FROM facilities
        WHERE facility_id = ? LIMIT 1;
    """, (facility_id,))
    row = cur.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Facility not found")

    (
      fid, name, group, tok, status, ca, ua, extras,
      legal_name, addr1, addr2, city, state, zip_, county,
      avg_dcs, short_beds, ltc_beds, outpatient_pt,
      emr, emr_other, pt_emr, orders, orders_other,
      raw_json
    ) = row

    if not tok or not t or t != tok:
        raise HTTPException(status_code=401, detail="Invalid intake token")

    # 2) Children (contacts/partners)
    cur.execute("SELECT type, name, email, phone, pref FROM facility_contacts WHERE facility_id = ? ORDER BY id ASC;", (fid,))
    contacts = [
        {"type": r[0] or "", "name": r[1] or "", "email": r[2] or "", "phone": r[3] or "", "pref": r[4] or ""}
        for r in cur.fetchall()
    ]

    cur.execute("SELECT type, name, ins_only, insurance FROM facility_partners WHERE facility_id = ? ORDER BY id ASC;", (fid,))
    partners = [
        {"type": r[0] or "", "name": r[1] or "", "ins_only": r[2] or "", "insurance": r[3] or ""}
        for r in cur.fetchall()
    ]

    # 3) Build response
    base = {
        "facility_id": fid,
        "facility_name": name or "",
        "corporate_group": group or "",
        "intake_status": status or "not-started",
        "created_at": ca, "updated_at": ua,
        "extras": extras or "",
        # intake fields saved on the facilities table
        "legal_name": legal_name or "",
        "address_line1": addr1 or "",
        "address_line2": addr2 or "",
        "city": city or "",
        "state": state or "",
        "zip": zip_ or "",
        "county": county or "",
        "avg_dcs": avg_dcs or "",
        "short_beds": short_beds or "",
        "ltc_beds": ltc_beds or "",
        "outpatient_pt": outpatient_pt or "",
        "emr": emr or "",
        "emr_other": emr_other or "",
        "pt_emr": pt_emr or "",
        "orders": orders or "",
        "orders_other": orders_other or "",
    }

    # 4) You may also store these as JSON in facilities.extras/raw_json; we return arrays directly
    return {"ok": True, "facility": base, "contacts": contacts, "community_partners": partners}



def _init_db():
    conn = _db()
    cur = conn.cursor()

    # --- qa table (mirrors qa.csv columns; add topics for back-compat) ---
    cur.execute("""
    CREATE TABLE IF NOT EXISTS qa (
        id TEXT PRIMARY KEY,
        section TEXT,
        topics TEXT,
        question TEXT NOT NULL,
        answer TEXT NOT NULL,
        tags TEXT
    );
    """)


    # NEW: facility quick facts table (row-per-fact)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS fac_facts (
      id           INTEGER PRIMARY KEY AUTOINCREMENT,
      facility_id  TEXT NOT NULL,
      fact_text    TEXT NOT NULL,
      tags         TEXT DEFAULT '',
      created_at   TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
      FOREIGN KEY (facility_id) REFERENCES facilities(facility_id) ON DELETE CASCADE
    );
    """)


    # --- client_info submissions (optional, stores latest snapshot per facility+field) ---
    # Not used yet; we'll use it in a later step if you choose the "upsert" path.
    cur.execute("""
    CREATE TABLE IF NOT EXISTS client_info (
        facility_name TEXT,
        field_name TEXT,
        value TEXT,
        updated_at TEXT,
        PRIMARY KEY (facility_name, field_name)
    );
    """)

    # --- abbreviations (aligns with your "abbr,meaning,notes") ---
    cur.execute("""
    CREATE TABLE IF NOT EXISTS abbreviations (
        abbr TEXT PRIMARY KEY,
        meaning TEXT,
        notes TEXT
    );
    """)

    # --- unified dictionary (general-purpose aliases/synonyms/abbreviations) ---
    cur.execute("""
    CREATE TABLE IF NOT EXISTS dictionary (
        key         TEXT PRIMARY KEY,
        canonical   TEXT NOT NULL,
        kind        TEXT DEFAULT 'abbr',
        notes       TEXT DEFAULT '',
        match_mode  TEXT DEFAULT 'exact'  -- 'exact' | 'word'
    );
    """)

    # ---- one-time migration from abbreviations -> dictionary (idempotent) ----
    cur.execute("""
    INSERT INTO dictionary(key, canonical, kind, notes, match_mode)
    SELECT abbr, meaning, 'abbr', COALESCE(notes,''), 'exact'
    FROM abbreviations
    WHERE abbr IS NOT NULL AND abbr <> ''
      AND NOT EXISTS (SELECT 1 FROM dictionary d WHERE d.key = abbreviations.abbr);
    """)



    # --- user_qa_log (extends with promoted, a_quality; matches latest CSV shape) ---
    cur.execute("""
    CREATE TABLE IF NOT EXISTS user_qa_log (
        id INTEGER PRIMARY KEY,
        ts TEXT,
        section TEXT,
        q TEXT,
        a TEXT,
        promoted INTEGER DEFAULT 0,
        a_quality TEXT DEFAULT ''
    );
    """)

    # --- facilities (core master record for each facility) ---
    cur.execute("""
    CREATE TABLE IF NOT EXISTS facilities (
        facility_id     TEXT PRIMARY KEY,          -- stable slug/ID
        facility_name   TEXT NOT NULL,             -- display name
        legal_name      TEXT,
        corporate_group TEXT,

        -- Address / location
        address_line1   TEXT,
        address_line2   TEXT,
        city            TEXT,
        state           TEXT,
        zip             TEXT,
        county          TEXT,

        -- Capacity / ops
        avg_dcs         TEXT,
        short_beds      TEXT,
        ltc_beds        TEXT,
        outpatient_pt   TEXT,

        -- Systems / workflows
        emr             TEXT,
        emr_other       TEXT,
        pt_emr          TEXT,
        orders          TEXT,
        orders_other    TEXT,

        -- Intake / meta
        intake_token    TEXT UNIQUE,
        intake_status   TEXT DEFAULT 'not-started',
        created_at      TEXT,
        updated_at      TEXT,
        extras          TEXT,
        raw_json        TEXT                       -- full payload snapshot
    );

    """)

    cur.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_facilities_token ON facilities(intake_token);")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_facilities_name ON facilities(facility_name);")

    # --- migrate missing facility columns (safe if DB already exists) ---
    _migrate_facilities_columns(cur)

    # --- Ask-GPT answer-ready views (drop + recreate for safety) ---
    cur.execute("DROP VIEW IF EXISTS v_facility_summary;")
    cur.execute("""
    CREATE VIEW v_facility_summary AS
    SELECT
      f.facility_id,
      f.facility_name,
      f.city,
      f.state,
      f.zip,
      COALESCE(NULLIF(f.emr,''), f.emr_other) AS emr,
      f.orders,
      f.orders_other,
      f.outpatient_pt,
      f.short_beds,
      f.ltc_beds,
      f.avg_dcs
    FROM facilities f;
    """)

    cur.execute("DROP VIEW IF EXISTS v_contacts_by_type;")
    cur.execute("""
    CREATE VIEW v_contacts_by_type AS
    SELECT facility_id, type, GROUP_CONCAT(name, ' | ') AS names
    FROM facility_contacts
    GROUP BY facility_id, type;
    """)

    cur.execute("DROP VIEW IF EXISTS v_partners_by_type;")
    cur.execute("""
    CREATE VIEW v_partners_by_type AS
    SELECT facility_id, type, GROUP_CONCAT(name, ' | ') AS names
    FROM facility_partners
    GROUP BY facility_id, type;
    """)

    conn.commit()
    conn.close()


def _migrate_facilities_columns(cur):
    # helper to “ADD COLUMN IF NOT EXISTS” in SQLite style
    cur.execute("PRAGMA table_info(facilities);")
    have = {row[1] for row in cur.fetchall()}
    def add(col, typ):
        if col not in have:
            cur.execute(f"ALTER TABLE facilities ADD COLUMN {col} {typ};")

    add("legal_name", "TEXT")
    add("address_line1", "TEXT")
    add("address_line2", "TEXT")
    add("city", "TEXT")
    add("state", "TEXT")
    add("zip", "TEXT")
    add("county", "TEXT")
    add("avg_dcs", "TEXT")
    add("short_beds", "TEXT")
    add("ltc_beds", "TEXT")
    add("outpatient_pt", "TEXT")
    add("emr", "TEXT")
    add("emr_other", "TEXT")
    add("pt_emr", "TEXT")
    add("orders", "TEXT")
    add("orders_other", "TEXT")
    add("raw_json", "TEXT")



    # --- child tables for repeating data ---

    # Additional Services (chips widget)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS facility_additional_services (
      id           INTEGER PRIMARY KEY AUTOINCREMENT,
      facility_id  TEXT NOT NULL,
      service      TEXT NOT NULL,
      UNIQUE(facility_id, service),
      FOREIGN KEY (facility_id) REFERENCES facilities(facility_id) ON DELETE CASCADE
    );
    """)

    # Insurance Plans (chips widget)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS facility_insurance_plans (
      id           INTEGER PRIMARY KEY AUTOINCREMENT,
      facility_id  TEXT NOT NULL,
      plan         TEXT NOT NULL,
      UNIQUE(facility_id, plan),
      FOREIGN KEY (facility_id) REFERENCES facilities(facility_id) ON DELETE CASCADE
    );
    """)

    # Contacts (new compact widget)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS facility_contacts (
      id           INTEGER PRIMARY KEY AUTOINCREMENT,
      facility_id  TEXT NOT NULL,
      type         TEXT NOT NULL,  -- Administrator | Social Services | Nursing | Therapy | Admissions | Physician/NP | Other Staff
      name         TEXT NOT NULL,
      email        TEXT,
      phone        TEXT,
      pref         TEXT,           -- Email | Call Office | Call Cell | Text Cell
      FOREIGN KEY (facility_id) REFERENCES facilities(facility_id) ON DELETE CASCADE
    );
    """)
    cur.execute("CREATE INDEX IF NOT EXISTS idx_contacts_fac ON facility_contacts(facility_id);")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_contacts_type ON facility_contacts(type);")

    # Community Partners (new compact widget)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS facility_partners (
      id           INTEGER PRIMARY KEY AUTOINCREMENT,
      facility_id  TEXT NOT NULL,
      type         TEXT NOT NULL,   -- Home Health | DME | Oxygen | Wound Vac | Tube Feeding Supplies | Other
      name         TEXT NOT NULL,   -- Agency name
      ins_only     TEXT,            -- Yes | No | ''
      insurance    TEXT,            -- free text if ins_only == Yes
      FOREIGN KEY (facility_id) REFERENCES facilities(facility_id) ON DELETE CASCADE
    );
    """)
    cur.execute("CREATE INDEX IF NOT EXISTS idx_partners_fac ON facility_partners(facility_id);")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_partners_type ON facility_partners(type);")



# Run once at startup so the file/tables exist before we hit endpoints
_init_db()


# [FN:normalize_abbrev_query] -----------------------------------------------
def normalize_abbrev_query(s: str) -> str:
    """
    Lightweight normalizer used for cache keys and fuzzy compares.
    Keeps it simple: lowercase + collapse whitespace + strip punctuation.
    If you later want to expand abbreviations via your Dictionary table,
    hook that logic in here (but keep it fast).
    """
    import re
    s = (s or "").strip().lower()
    # Replace any non-alphanumeric with space, collapse spaces
    s = re.sub(r"[^a-z0-9]+", " ", s)
    return " ".join(s.split())

# [FN:_detect_facility_from_question] ----------------------------------------
def _detect_facility_from_question(q: str) -> str | None:
    """
    Extract a likely facility identifier from a natural question.
    Strategy:
      1) If quoted text exists, treat first quoted segment as the facility.
      2) Look for 'for <name>' or 'at <name>'.
      3) Fuzzy match against facilities table names/ids.
    Returns a string to feed into _facility_details_answer_from_db, or None.
    """
    import re
    text = (q or "").strip()
    if not text:
        return None

    # 1) Quoted capture: “...”, "...", '...'
    m = re.search(r"[\"“”']([^\"“”']+)[\"“”']", text)
    if m:
        cand = m.group(1).strip()
        if cand:
            return cand

    # 2) Phrases like: facility details for X / details for X / at X
    m = re.search(r"\b(?:for|at)\s+([a-z0-9][a-z0-9\s\-\&\.\']{1,60})$", text, flags=re.I)
    if m:
        cand = m.group(1).strip(" .-'\"")
        if cand:
            return cand

    # 3) Fuzzy DB match on facilities table
    try:
        conn = _db()
        cur  = conn.cursor()
        cur.execute("SELECT facility_id, facility_name FROM facilities;")
        rows = cur.fetchall()
        norm_text = re.sub(r"[^a-z0-9]+", " ", text.lower()).split()
        best = None
        best_score = 0
        for fid, name in rows:
            tokens = re.sub(r"[^a-z0-9]+", " ", (name or "").lower()).split()
            # simple token overlap score
            score = sum(1 for t in tokens if t in norm_text)
            if (fid or "").lower() in norm_text:
                score += 2
            if score > best_score:
                best = name or fid
                best_score = score
        return best
    except Exception:
        return None

# [FN:_contact_answer_from_db] ---------------------------------------------
def _contact_answer_from_db(facility_name: str, role: str) -> tuple[str | None, dict]:
    """
    Returns (markdown, meta) for a contact role at a facility from SQLite.
    """
    if not facility_name or not role:
        return None, {}

    # Resolve facility_id by name or id (simple match)
    conn = _db()
    cur  = conn.cursor()
    cur.execute("SELECT facility_id, facility_name FROM facilities;")
    rows = cur.fetchall()
    fid = None
    for f_id, f_name in rows:
        if facility_name.strip().lower() in {(f_id or '').lower(), (f_name or '').lower()}:
            fid = f_id
            break
    if not fid:
        # try fuzzy: token overlap
        txt = re.sub(r"[^a-z0-9]+", " ", facility_name.lower())
        best, best_score = None, 0
        for f_id, f_name in rows:
            score = sum(1 for t in re.sub(r"[^a-z0-9]+"," ", (f_name or '').lower()).split() if t in txt.split())
            if (f_id or "").lower() in txt.split():
                score += 2
            if score > best_score:
                best, best_score = f_id, score
        fid = best

    if not fid:
        return None, {}

    # Pull the matching contact
    cur.execute("""
        SELECT type, name, email, phone, pref
        FROM facility_contacts
        WHERE facility_id = ?
          AND LOWER(type) = LOWER(?)
        ORDER BY id ASC
        LIMIT 1;
    """, (fid, role))
    row = cur.fetchone()
    if not row:
        # allow loose contains match
        cur.execute("""
            SELECT type, name, email, phone, pref
            FROM facility_contacts
            WHERE facility_id = ?
              AND LOWER(type) LIKE '%' || LOWER(?) || '%'
            ORDER BY id ASC
            LIMIT 1;
        """, (fid, role))
        row = cur.fetchone()

    if not row:
        return None, {}

    ctype, name, email, phone, pref = row
    md = f"""**{ctype} for {facility_name.title()}**  
- Name: {name or '—'}  
- Email: {email or '—'}  
- Phone: {phone or '—'}  
- Contact preference: {pref or '—'}"""
    return md.strip(), {"facility_id": fid, "role": role}


# [FN:_facility_details_answer_from_db] ----------------------------------------
def _facility_details_answer_from_db(fac_query: str) -> tuple[str, dict]:
    """
    Build a rich Facility Details answer directly from SQLite tables:
      - facilities
      - facility_contacts
      - facility_partners
      - facility_facts (if present)
    Matching tries facility_id first, then facility_name (case-insensitive, fuzzy).
    Returns: (markdown_answer, meta_used)  meta_used contains {'facility_id': ...}
    """
    import sqlite3, re
    q = (fac_query or "").strip()
    if not q:
        return "", {}

    def _canon(s: str) -> str:
        s = (s or "").strip().lower()
        s = re.sub(r"[^a-z0-9]+", " ", s)
        return " ".join(s.split())

    c_q = _canon(q)

    conn = _db()
    cur  = conn.cursor()

    # 1) try exact id match
    cur.execute("SELECT facility_id, facility_name FROM facilities WHERE facility_id = ? LIMIT 1;", (q,))
    row = cur.fetchone()

    # 2) else fuzzy on facility_name
    if not row:
        cur.execute("""
            SELECT facility_id, facility_name
            FROM facilities
            ORDER BY facility_name COLLATE NOCASE ASC;
        """)
        rows = cur.fetchall()
        best = None
        for fid, name in rows:
            if _canon(name) == c_q:
                best = (fid, name); break
            if c_q and _canon(name).startswith(c_q):
                best = best or (fid, name)
        if not best and rows:
            # very loose contains
            for fid, name in rows:
                if c_q and c_q in _canon(name):
                    best = (fid, name); break
        row = best

    if not row:
        return "", {}

    fid, fac_name = row

    # 3) pull facility core
    cur.execute("""
        SELECT facility_id, facility_name, corporate_group, intake_status,
               legal_name, address_line1, address_line2, city, state, zip, county,
               avg_dcs, short_beds, ltc_beds, outpatient_pt,
               emr, emr_other, pt_emr, orders, orders_other
        FROM facilities WHERE facility_id = ? LIMIT 1;
    """, (fid,))
    core = cur.fetchone()
    if not core:
        return "", {}

    (_fid, name, corp, status,
     legal, a1, a2, city, st, z, county,
     avg_dcs, short_beds, ltc_beds, outpt,
     emr, emr_other, pt_emr, orders, orders_other) = core

    # 4) contacts
    cur.execute("""
        SELECT type, name, email, phone, pref
        FROM facility_contacts
        WHERE facility_id = ?
        ORDER BY id ASC;
    """, (fid,))
    contacts = [{"type": r[0] or "", "name": r[1] or "", "email": r[2] or "", "phone": r[3] or "", "pref": r[4] or ""} for r in cur.fetchall()]

    # 5) community partners
    try:
        cur.execute("""
            SELECT type, name, ins_only, insurance
            FROM facility_partners
            WHERE facility_id = ?
            ORDER BY id ASC;
        """, (fid,))
        partners = [{"type": r[0] or "", "name": r[1] or "", "ins_only": r[2] or "", "insurance": r[3] or ""} for r in cur.fetchall()]
    except Exception:
        partners = []

    # 6) facts (if the table exists)
    facts = []
    try:
        cur.execute("SELECT statement, tag FROM facility_facts WHERE facility_id = ? ORDER BY created_at DESC, id DESC;", (fid,))
        facts = [{"statement": r[0] or "", "tag": r[1] or ""} for r in cur.fetchall()]
    except Exception:
        pass


    # 6.1) additional services (if the table exists)
    additional_services = []
    try:
        # common schemas seen: (facility_id, service) | (facility_id, type) | (facility_id, name)
        cur.execute("PRAGMA table_info(facility_additional_services);")
        cols = [c[1].lower() for c in cur.fetchall()]
        svc_col = next((c for c in ("service","type","name","label") if c in cols), None)
        if svc_col:
            cur.execute(f"""
                SELECT {svc_col}
                FROM facility_additional_services
                WHERE facility_id = ?
                ORDER BY 1 COLLATE NOCASE ASC;
            """, (fid,))
            additional_services = [ (r[0] or "").strip() for r in cur.fetchall() if (r[0] or "").strip() ]
    except Exception:
        pass

    # 6.2) insurance plans (if the table exists)
    ins_plans = []
    try:
        # common schemas: (facility_id, plan) | (facility_id, plan_name) | (facility_id, name)
        cur.execute("PRAGMA table_info(insurance_plans);")
        cols = [c[1].lower() for c in cur.fetchall()]
        plan_col = next((c for c in ("plan","plan_name","name") if c in cols), None)
        if plan_col:
            cur.execute(f"""
                SELECT {plan_col}
                FROM insurance_plans
                WHERE facility_id = ?
                ORDER BY 1 COLLATE NOCASE ASC;
            """, (fid,))
            ins_plans = [ (r[0] or "").strip() for r in cur.fetchall() if (r[0] or "").strip() ]
    except Exception:
        pass




    # 7) render markdown
    def _row(label, val):
        val = (str(val or "").strip())
        return f"- **{label}:** {val}" if val else ""

    lines = []
    lines.append(f"### {name or fid} — Facility Details\n")
    lines.extend([x for x in [
        _row("Facility ID", fid),
        _row("Corporate Group", corp),
        _row("Intake Status", status),
        _row("Legal Name", legal),
        _row("Address", " ".join([x for x in [a1, a2] if x])),
        _row("City/State/Zip", " ".join([x for x in [city, st, z] if x])),
        _row("County", county),
        _row("Average DCs/Month", avg_dcs),
        _row("Beds (Short / LTC)", " / ".join([str(short_beds or ""), str(ltc_beds or "")]).strip(" /")),
        _row("Outpatient PT", outpt),
        _row("EMR", (emr_other if (emr or "").lower()=="other" else emr)),
        _row("PT EMR", pt_emr),
        _row("Orders System", (orders_other if (orders or "").lower()=="other" else orders)),
    ] if x])

    # contacts
    if contacts:
        lines.append("\n**Key Contacts**")
        for c in contacts:
            line = " - " + " | ".join([p for p in [
               f"_{c['type']}_".strip("_ "),
               c["name"], c["email"], c["phone"],
               f"pref: {c['pref']}" if c["pref"] else ""
            ] if p])
            lines.append(line)

    # partners
    if partners:
        lines.append("\n**Community Partners**")
        for p in partners:
            tag = ""
            if p["ins_only"] and p["insurance"]:
                tag = f" (only for {p['insurance']})"
            elif p["ins_only"]:
                tag = " (insurance-specific)"
            lines.append(f" - _{p['type']}_ — {p['name']}{tag}")

    # additional services
    if additional_services:
        lines.append("\n**Additional Services**")
        for s in additional_services:
            lines.append(f" - {s}")

    # insurance plans
    if ins_plans:
        lines.append("\n**Insurance Plans**")
        for p in ins_plans:
            lines.append(f" - {p}")

    # facts
    if facts:
        lines.append("\n**Quick Facts**")
        for f in facts[:12]:
            label = f" [{f['tag']}]" if f.get("tag") else ""
            lines.append(f" - {f['statement']}{label}")

    md = "\n".join(lines).strip()
    return (md, {"facility_id": fid})



# ------------------------------------------------------------------------------
# [S5.2] SQLITE helpers (read-only for Step 2)
# ------------------------------------------------------------------------------
def db_fetchall(query, params=()):
    conn = _db()
    cur  = conn.cursor()
    cur.execute(query, params)
    cols = [c[0] for c in cur.description]
    out = [dict(zip(cols, row)) for row in cur.fetchall()]
    conn.close()
    return out

def qa_df_from_sqlite():
    rows = db_fetchall("""
        SELECT id, section, topics, question, answer, tags
        FROM qa
        ORDER BY id
    """)

    import pandas as pd
    import numpy as np

    df = pd.DataFrame(rows)
    if df is None or df.empty:
        # Ensure expected columns exist even if empty
        df = pd.DataFrame(columns=["id","section","topics","question","answer","tags"])

    # Make sure all expected columns exist
    for col in ["id","section","topics","question","answer","tags"]:
        if col not in df.columns:
            df[col] = ""

    # Normalize to strings (avoids None/NaN surprises)
    for col in ["section","topics","question","answer","tags"]:
        df[col] = df[col].astype(str).fillna("")

    # Safely synthesize topics from section if topics blank
    has_topics = df["topics"].str.strip() != ""
    df.loc[~has_topics, "topics"] = df.loc[~has_topics, "section"].str.strip()

    return df


    # Defensive: ensure expected columns exist even on older DBs
    for col in ["id", "section", "topics", "question", "answer", "tags"]:
        if col not in df.columns:
            df[col] = ""

    # Back-compat: synthesize topics from section when missing
    def _mk_topics(row):
        t = (row.get("topics") or "").strip()
        if t:
            return t
        s = (row.get("section") or "").strip()
        return s if s else ""
    df["topics"] = df.apply(_mk_topics, axis=1)

    # If SQLite tags are blank or missing for some rows, try to backfill from CSV by id
    try:
        csv_df = load_csv_df()  # already normalizes cols
        if not csv_df.empty:
            csv_tags = csv_df[["id", "tags"]].copy()
            csv_tags["id"] = csv_tags["id"].astype(str)
            df["id"] = df["id"].astype(str)
            df = df.merge(csv_tags, on="id", how="left", suffixes=("", "_csv"))
            # Prefer SQLite tags when present; otherwise take CSV tags
            df["tags"] = df["tags"].where(df["tags"].astype(str).str.strip() != "", df["tags_csv"].fillna(""))
            df = df.drop(columns=[c for c in ["tags_csv"] if c in df.columns])
    except Exception:
        # If anything goes wrong, we still return a df with a tags column
        pass

    return df


def abv_rows_from_sqlite():
    return db_fetchall("""
        SELECT abbr, meaning, notes
        FROM abbreviations
        ORDER BY abbr
    """)

def log_rows_from_sqlite(limit=1000):
    return db_fetchall(f"""
        SELECT id, ts, section, q, a, promoted, a_quality
        FROM user_qa_log
        ORDER BY id DESC
        LIMIT {int(limit)}
    """)


# ------------------------------------------------------------------------------
# [S5.3] Unified loaders (source = SQLite if USE_SQLITE_READ else CSV)
# ------------------------------------------------------------------------------
def load_qa_df_unified():
    if USE_SQLITE_READ:
        df = qa_df_from_sqlite()
    else:
        df = load_csv_df()

    # Final safety: ensure columns exist and are strings
    for col in ["tags","topics","section","question","answer"]:
        if col not in df.columns:
            df[col] = ""
        df[col] = df[col].astype(str).fillna("")

    # If topics still blank, mirror section
    df.loc[df["topics"].str.strip() == "", "topics"] = df["section"].str.strip()

    return df



def load_abbrev_unified():
    if USE_SQLITE_READ:
        return abv_rows_from_sqlite()
    return _read_abv_rows()

def load_userlog_unified(limit=1000):
    if USE_SQLITE_READ:
        return log_rows_from_sqlite(limit=limit)
    return _read_log_rows()


# ------------------------------------------------------------------------------
# [S5.4] SQLITE upsert helpers (called after CSV writes in Step 3)
# ------------------------------------------------------------------------------
def _exec_write(sql, params=()):
    conn = _db()
    cur  = conn.cursor()
    cur.execute(sql, params)
    conn.commit()
    conn.close()

def upsert_qa_sqlite(row: dict):
    _exec_write("""
        INSERT INTO qa (id, section, topics, question, answer, tags)
        VALUES (?, ?, ?, ?, ?, ?)
        ON CONFLICT(id) DO UPDATE SET
          section = excluded.section,
          topics  = excluded.topics,
          question= excluded.question,
          answer  = excluded.answer,
          tags    = excluded.tags
    """, (
        str(row.get("id","")).strip(),
        str(row.get("section","")).strip(),
        str(row.get("topics","")).strip(),
        str(row.get("question","")).strip(),
        str(row.get("answer","")).strip(),
        str(row.get("tags","")).strip(),
    ))

def delete_qa_sqlite(qid: str):
    _exec_write("DELETE FROM qa WHERE id = ?", (str(qid),))

def upsert_log_sqlite(row: dict):
    _exec_write("""
        INSERT OR REPLACE INTO user_qa_log (id, ts, section, q, a, promoted, a_quality)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (
        int(row.get("id", 0) or 0),
        str(row.get("ts","")),
        str(row.get("section","")),
        str(row.get("q","")),
        str(row.get("a","")),
        1 if str(row.get("promoted","")).strip() in ("1","true","True") else 0,
        str(row.get("a_quality","")),
    ))


def upsert_dictionary_sqlite(row: dict):
    _exec_write("""
        INSERT INTO dictionary (key, canonical, kind, notes, match_mode)
        VALUES (?, ?, ?, ?, ?)
        ON CONFLICT(key) DO UPDATE SET
          canonical = excluded.canonical,
          kind      = excluded.kind,
          notes     = excluded.notes,
          match_mode= excluded.match_mode
    """, (
        str(row.get("key","")).strip(),
        str(row.get("canonical","")).strip(),
        str(row.get("kind","abbr")).strip(),
        str(row.get("notes","")).strip(),
        str(row.get("match_mode","exact")).strip().lower() or "exact",
    ))
    clear_dictionary_caches()


def upsert_abbrev_sqlite(row: dict):
    _exec_write("""
        INSERT INTO abbreviations (abbr, meaning, notes)
        VALUES (?, ?, ?)
        ON CONFLICT(abbr) DO UPDATE SET
          meaning = excluded.meaning,
          notes   = excluded.notes
    """, (
        str(row.get("abbr","")).strip(),
        str(row.get("meaning","")).strip(),
        str(row.get("notes","")).strip(),
    ))

def _next_log_id_sqlite():
    row = db_fetchall("SELECT COALESCE(MAX(id), 0) AS mx FROM user_qa_log")[0]
    return int(row["mx"]) + 1


# ==============================================================================
# [S6] MODELS, PROMPTS & ASK PAYLOAD
# ==============================================================================
# =========================
# Models & Prompts
# =========================
class AskPayload(BaseModel):
    question: str
    top_k: int = 4
    section: str | None = None          # UI may send this…
    section_hint: str | None = None     # …or this (from the Copilot page)
    debug: bool = False                 # NEW: return scoring details when True



SYSTEM_PROMPT = (
    "You are the Evolv Health training copilot.\n"
    "- Answer STRICTLY from the provided context when possible.\n"
    "- If the answer appears ANYWHERE in the context, you MUST give it; do NOT say it's missing.\n"
    "- Keep answers concise and accurate.\n"
    "- DO NOT include source IDs or citations in the answer text. The client will display sources separately."
)





# ==============================================================================
# [S7] HELPERS — Admin check, CSV I/O, embed, MMR, etc.
# ==============================================================================
# =========================
# Helpers
# =========================

# [FN:require_admin] ----------------------------------------------------
def require_admin(req: Request):
    """Accept token via header OR query OR form to avoid CORS preflight headaches."""
    # Allow CORS preflight to pass without a token
    if req.method == "OPTIONS":
        return True
    token = (
        req.headers.get("x-admin-token")
        or req.query_params.get("token")
        or ""
    )
    if not ADMIN_TOKEN or token != ADMIN_TOKEN:
        raise HTTPException(status_code=401, detail="Unauthorized")

# [FN:_assert_admin] ----------------------------------------------------
def _assert_admin(token: str):
    if not ADMIN_TOKEN or token != ADMIN_TOKEN:
        raise HTTPException(status_code=401, detail="Unauthorized")


# [FN:load_csv_df] ----------------------------------------------------
def load_csv_df() -> pd.DataFrame:
    if not os.path.exists(CSV_PATH):
        os.makedirs(os.path.dirname(CSV_PATH), exist_ok=True)
        with open(CSV_PATH, "w", newline="", encoding="utf-8") as f:
            # new header includes topics; keep section for back-compat
            f.write("id,section,topics,question,answer,tags\n")
    df = pd.read_csv(CSV_PATH).fillna("")

    # Ensure columns exist (back-compat with older CSVs)
    for col in ["id", "section", "topics", "question", "answer", "tags"]:
        if col not in df.columns:
            df[col] = ""

    # Back-compat: synthesize topics from section when topics empty
    def _mk_topics(row):
        t = (row.get("topics") or "").strip()
        if t:
            return t
        s = (row.get("section") or "").strip()
        return s if s else ""

    tmp = df.apply(_mk_topics, axis=1)
    # If _mk_topics returns a Series/dict per row, 'tmp' becomes a DataFrame.
    # If it returns a scalar/list per row, 'tmp' is a Series.
    if isinstance(tmp, pd.DataFrame):
        # Pick just the 'topics' piece for the column assignment.
        if "topics" in tmp.columns:
            df["topics"] = tmp["topics"]
        else:
            # Fallback: turn whole row result into a topics list/string
            df["topics"] = tmp.apply(lambda x: x.get("topics", []) if hasattr(x, "get") else x)
        # Optional: if _mk_topics also normalized 'section', keep it.
        if "section" in tmp.columns:
            df["section"] = tmp["section"]
    else:
        # Simple case: _mk_topics returned a scalar/list per row
        df["topics"] = tmp

    df["id"] = df["id"].astype(str)
    return df


# [FN:save_csv_df] ----------------------------------------------------
def save_csv_df(df: pd.DataFrame):
    tmp = CSV_PATH + ".tmp"
    df.to_csv(tmp, index=False, quoting=csv.QUOTE_MINIMAL, encoding="utf-8")
    os.replace(tmp, CSV_PATH)  # atomic on Windows


# [FN:embed_query] ----------------------------------------------------
def embed_query(q: str):
    resp = client.embeddings.create(model=EMBED_MODEL, input=[q])
    return resp.data[0].embedding


# [FN:cosine] ----------------------------------------------------
def cosine(u, v):
    # u, v: dense vectors (lists of floats)
    if not u or not v or len(u) != len(v):
        return 0.0
    num = sum(a*b for a, b in zip(u, v))
    du = math.sqrt(sum(a*a for a in u))
    dv = math.sqrt(sum(b*b for b in v))
    return (num / (du*dv)) if du and dv else 0.0


# [FN:mmr_select] ----------------------------------------------------
def mmr_select(query_vec, candidates, lambda_=0.65, topn=6):
    """
    MMR: select topn items maximizing relevance - redundancy.
    candidates: list of dicts with 'embedding' (list), 'score' (float baseline), and other fields
    """
    selected = []
    remaining = candidates[:]
    while remaining and len(selected) < topn:
        best, best_val = None, -1e9
        for c in remaining:
            rel = cosine(query_vec, c.get("embedding") or [])
            if math.isnan(rel): rel = 0.0
            # redundancy = max sim to already selected
            if selected:
                red = max(cosine(c.get("embedding") or [], s.get("embedding") or []) for s in selected)
            else:
                red = 0.0
            val = lambda_ * rel - (1 - lambda_) * red
            if val > best_val:
                best, best_val = c, val
        selected.append(best)
        remaining.remove(best)
    return selected


@lru_cache(maxsize=1024)
# [FN:generate_query_expansions] ----------------------------------------------------
def generate_query_expansions(q: str, n: int = 4) -> list[str]:
    """
    Use the model to produce paraphrases/expansions (abbreviations, synonyms, formal names).
    This is general-purpose; no client-specific rules.
    """
    prompt = (
        "Rewrite the user question into diverse alternative phrasings that could retrieve the same facts. "
        "Include abbreviations and expanded facility names if applicable. "
        "Return one variant per line, no numbering.\n\n"
        f"User question: {q}"
    )
    resp = client.responses.create(
        model="gpt-5-mini",
        input=[{"role": "user", "content": prompt}],
        max_output_tokens=200,
    )
    text = (resp.output_text or "").strip()
    lines = [ln.strip("-• \t") for ln in text.splitlines() if ln.strip()]
    # dedupe and cap to n
    out = list(OrderedDict.fromkeys(lines))[:n]
    # Always include the original
    if q not in out:
        out.insert(0, q)
    return out


# [FN:retrieve] ----------------------------------------------------
def retrieve(q: str, k: int, section: str | None = None):
    # dictionary-aware query text
    q = normalize_dictionary_query(q)
    k = max(1, min(k, 12))

    # 1) multi-query expansion (fewer when section is given)
    n_variants = 2 if section else 4
    variants = generate_query_expansions(q, n=n_variants)
    if q not in variants:
        variants = [q] + variants

    # 2) batch embed variants (first is q)
    emb = client.embeddings.create(model=EMBED_MODEL, input=variants)
    vmap = {txt: d.embedding for txt, d in zip(variants, emb.data)}
    qvec = vmap[q]

    # 3) accumulate candidates
    candidates = []

    def add_results(res, kind, w):
        if not res or not res.get("ids") or not res["ids"][0]:
            return
        ids   = res["ids"][0]
        docs  = res["documents"][0]
        metas = res.get("metadatas", [[]])[0] if res.get("metadatas") else [{} for _ in ids]
        dists = res.get("distances", [[]])[0] if res.get("distances") else [1.0 for _ in ids]
        embs  = res.get("embeddings", [[]])[0] if res.get("embeddings") else [[] for _ in ids]

        for i in range(len(ids)):
            dist  = float(dists[i]) if i < len(dists) else 1.0
            base_score = w * (1.0 / (1.0 + dist))
            emb_vec = None
            if i < len(embs):
                ev = embs[i]
                # coerce numpy arrays to plain lists
                if hasattr(ev, "tolist"):
                    emb_vec = ev.tolist()
                elif isinstance(ev, (list, tuple)):
                    emb_vec = list(ev)
                else:
                    emb_vec = None

            candidates.append({
                "kind": kind,
                "id": str(ids[i]),
                "text": docs[i],
                "meta": metas[i] if i < len(metas) else {},
                "dist": dist,
                "score": base_score,
                "embedding": emb_vec,
            })

    # Query across variants; pull a bit more than k each time
    per_variant = min(6, max(4, k))
    for v in variants:
        vvec = vmap[v]
        qa_res   = _safe_query(
            collection, vvec, per_variant,
            include=["documents","metadatas","distances","embeddings"],
            rebuild_fn=rebuild_index_from_csv
        )

        # NEW: for certain operational topics, prefer Q&A and skip docs
        topics_lower = [t.strip().lower() for t in (section or "").split(",") if t.strip()]
        QA_ONLY_TOPICS = {
            "sensys", "ur calls", "dme confirmation", "hh start of care", "closing case call"
        }
        skip_docs = any(t in QA_ONLY_TOPICS for t in topics_lower)

        if skip_docs:
            docs_res = {"ids":[[]], "documents":[[]], "metadatas":[[]], "distances":[[]], "embeddings":[[]]}
        else:
            docs_res = _safe_query(
                docs_collection, vvec, per_variant,
                include=["documents","metadatas","distances","embeddings"],
                rebuild_fn=rebuild_docs_index_from_folder
            )

        add_results(qa_res,   "qa",  QA_WEIGHT)
        add_results(docs_res, "doc", DOCS_WEIGHT)


    # 4) de-dupe by (kind,id) keeping the best base score
    uniq = {}
    for c in candidates:
        key = (c["kind"], c["id"])
        if key not in uniq or c["score"] > uniq[key]["score"]:
            uniq[key] = c
    candidates = list(uniq.values())

    # 5) If embeddings missing (older Chroma), embed candidate texts now for MMR
    def _is_missing(v):
        if v is None:
            return True
        try:
            import numpy as np
            if isinstance(v, np.ndarray):
                return v.size == 0
        except Exception:
            pass
        try:
            return len(v) == 0
        except Exception:
            return False

    need_vecs = any(_is_missing(c.get("embedding")) for c in candidates)
    if need_vecs:
        embs = client.embeddings.create(model=EMBED_MODEL, input=[c["text"] for c in candidates])
        vecs = [d.embedding for d in embs.data]
        for c, v in zip(candidates, vecs):
            c["embedding"] = v

    # ----- Facility-aware filtering (HARD when a facility is named) ------------
    detected_facility = _detect_facility_from_question(q)
    if detected_facility:
        candidates = [
            c for c in candidates
            if c["kind"] != "qa" or _match_facility_in_meta(c.get("meta", {}), detected_facility)
        ]
    # ----------------------------------------------------------------------------

    # ----- Topic-aware filtering / boosting (replaces old 'section' equality) ---
    if section:
        # section may now be a comma list of topics
        req_topics = [t.strip().lower() for t in section.split(",") if t.strip()]

        def _row_topics(meta: dict) -> list[str]:
            # prefer 'topics'; fallback to single 'section'
            t = str(meta.get("topics", "") or meta.get("section", "") or "")
            return [x.strip().lower() for x in t.split(",") if x.strip()]

        wanted = set(req_topics)
        # Strict keep if ANY overlap
        filtered = [
            c for c in candidates
            if c["kind"] == "qa" and (set(_row_topics(c.get("meta", {}))) & wanted)
        ]
        if filtered:
            candidates = filtered
        else:
            # If no strict matches, gently boost items overlapping by topic
            for c in candidates:
                if c["kind"] != "qa":
                    continue
                rowt = set(_row_topics(c.get("meta", {})))
                if rowt & wanted:
                    c["score"] += 0.25
    # ---------------------------------------------------------------------------

    # ---- NEW: Tag-overlap boost (generic, helps "staff", "admissions", etc.) ---
    try:
        q_tokens = _tokenize(q)
        for c in candidates:
            if c["kind"] != "qa":
                continue
            tags = str(c.get("meta", {}).get("tags", "") or "")
            if not tags:
                continue
            boost = _tag_overlap_score(q_tokens, tags)
            if boost:
                # Weighted modestly so it nudges rather than overwhelms
                c["score"] += 0.5 * boost
    except Exception as e:
        print("tag boost skipped:", e)
    # ---------------------------------------------------------------------------

    # ---- Deterministic component scores (KW + TAG) -----------------------------
    try:
        q_tokens = _tokenize(q)
    except Exception:
        q_tokens = set()

    def _kw_overlap(meta: dict) -> float:
        blob = " ".join([
            str(meta.get("question","") or ""),
            str(meta.get("answer","")   or ""),
            str(meta.get("tags","")     or ""),
        ])
        toks = _tokenize(blob)
        return len(q_tokens & toks) / max(1.0, len(q_tokens))

    TAG_W = 0.6
    KW_W  = 0.4

    for c in candidates:
        c.setdefault("score", 0.0)
        c_id = str(c.get("id","") or "")
        c["_id_sort"] = c_id  # for stable tie-break

        if c["kind"] != "qa":
            c["_kw"] = 0.0
            c["_tagb"] = 0.0
            c["_pre"] = float(c["score"])
            continue

        meta = c.get("meta", {}) or {}
        kw   = _kw_overlap(meta)
        tagb = _tag_overlap_score(q_tokens, str(meta.get("tags","") or ""))
        pre  = KW_W*kw + TAG_W*tagb

        c["_kw"]  = kw
        c["_tagb"] = tagb
        c["_pre"]  = pre
        c["score"] = float(c["score"]) + pre

    # Stable pre-sort to kill jitter before MMR
    candidates.sort(
        key=lambda x: (x["score"], x.get("_pre",0.0), x.get("_kw",0.0), x.get("_tagb",0.0), x.get("_id_sort","")),
        reverse=True
    )

    # 6) MMR re-rank, then take top-k (keep deterministic)
    mmr = mmr_select(qvec, candidates, lambda_=0.65, topn=k)
    mmr.sort(
        key=lambda x: (x["score"], x.get("_pre",0.0), x.get("_kw",0.0), x.get("_tagb",0.0), x.get("_id_sort","")),
        reverse=True
    )
    return mmr


# ==============================================================================
# [S7.2] HELPERS — Intake templating & incremental index upsert
# ==============================================================================
import os as _os_s72, csv as _csv_s72, re as _re_s72
import pandas as _pd_s72

try:
    from slugify import slugify  # pip install python-slugify
except Exception:
    def slugify(x: str) -> str:
        return _re_s72.sub(r'[^a-z0-9]+', '-', str(x).strip().lower()).strip('-')

# Reuse existing OpenAI client & EMBED_MODEL
try:
    client  # noqa: F821
except NameError:
    from openai import OpenAI
    client = OpenAI()
try:
    EMBED_MODEL  # noqa: F821
except NameError:
    EMBED_MODEL = "text-embedding-3-large"

# Reuse Chroma setup
import chromadb as _chroma_s72
VECTOR_PATH = "vectorstore"
QA_COLLECTION = "qa_main"

CLIENT_SHEET_PATH = _os_s72.path.join("data", "Client_Info_Spreadsheet.csv")
QA_CSV_PATH       = _os_s72.path.join("data", "qa.csv")

# ... (all [S7.2] functions exactly as in your last working copy) ...
# Kept verbatim from your current file to avoid drift:
# _best_facility_name, _render_question (Row-2 templating), _make_qid,
# _embed_texts, _chroma, _upsert_into_chroma, _upsert_into_qa_csv,
# _is_blank, _clean_str, _generate_qas_for_submission, _build_tags_from_payload
#  — these appear later again under [S17]; we keep this [S7.2] subset
#  for intake-submit path and incremental index upserts.


# ==============================================================================
# [S7.3b] HELPERS — facility + category scoping for List/Group
# ==============================================================================
import re, pandas as pd, string

_TAG_SPLIT_RE = re.compile(r"[,\|;]")

def _split_tags(s: str) -> list[str]:
    return [t.strip() for t in _TAG_SPLIT_RE.split(str(s or "")) if t.strip()]

def _tags_normalized(s: str) -> list[str]:
    return [t.lower() for t in _split_tags(s)]

def _canon(s: str) -> str:
    """Lowercase, trim, collapse spaces & punctuation for fuzzy match."""
    if s is None:
        return ""
    t = str(s).lower().strip()
    t = t.translate(str.maketrans("", "", string.punctuation))
    t = re.sub(r"\s+", " ", t)
    return t

def _rows_from_qa_csv() -> list[dict]:
    rows = []
    try:
        df = load_csv_df()  # must include: question, answer, tags
        cols = {c.lower(): c for c in df.columns}
        qcol, acol, tcol = cols.get("question"), cols.get("answer"), cols.get("tags")
        if qcol and acol:
            for _, r in df.iterrows():
                rows.append({
                    "question": str(r.get(qcol, "") or ""),
                    "answer":   str(r.get(acol, "") or ""),
                    "tags":     str(r.get(tcol, "") or "")
                })
    except Exception:
        pass
    return rows

def _rows_from_client_info() -> list[dict]:
    """Optional: only contributes if your client sheet has question/answer (and maybe tags)."""
    rows = []
    try:
        cdf = pd.read_csv(CLIENT_SHEET_PATH)
        cols = {c.lower(): c for c in cdf.columns}
        qcol, acol, tcol = cols.get("question"), cols.get("answer"), cols.get("tags")
        if qcol and acol:
            for _, r in cdf.iterrows():
                rows.append({
                    "question": str(r.get(qcol, "") or ""),
                    "answer":   str(r.get(acol, "") or ""),
                    "tags":     str(r.get(tcol, "") or "")
                })
    except Exception:
        pass
    return rows

def _list_unique_facilities_from_sources() -> list[str]:
    names = []
    # qa.csv via f- tags
    for r in _rows_from_qa_csv():
        for t in _tags_normalized(r.get("tags", "")):
            if t.startswith("f-"):
                name = t[2:].strip()
                if name:
                    names.append(name)
    # client sheet: tags and/or facility_name
    try:
        cdf = pd.read_csv(CLIENT_SHEET_PATH)
        if "tags" in cdf.columns:
            for tags in cdf["tags"].fillna(""):
                for t in _tags_normalized(tags):
                    if t.startswith("f-"):
                        name = t[2:].strip()
                        if name:
                            names.append(name)
        if "facility_name" in cdf.columns:
            for v in cdf["facility_name"].fillna(""):
                v = str(v).strip()
                if v:
                    names.append(v)
    except Exception:
        pass
    # de-dupe, keep order
    seen, uniq = set(), []
    for n in names:
        if n and n not in seen:
            seen.add(n)
            uniq.append(n)
    return uniq

# --- question parsing ---
_LIST_ALL_RE = re.compile(r"\b(list|show)\s+(all|every)\s+([A-Za-z0-9 _\-]+)", re.IGNORECASE)

def _infer_category_from_question(q: str) -> str | None:
    m = _LIST_ALL_RE.search(q or "")
    if not m:
        return None
    cat = (m.group(3) or "").strip()
    if not cat:
        return None
    # light singularization
    if len(cat) > 3 and cat.lower().endswith("s"):
        cat = cat[:-1].strip()
    return cat.lower()

def _infer_facility_from_question(q: str) -> str | None:
    """
    Try to pick a facility name mentioned in the question by fuzzy matching against
    the known facilities list (from tags + client sheet).
    """
    cq = _canon(q)
    if not cq:
        return None
    facilities = _list_unique_facilities_from_sources()
    facilities_sorted = sorted(facilities, key=lambda x: len(x), reverse=True)
    for name in facilities_sorted:
        if _canon(name) and _canon(name) in cq:
            return name  # pretty/original casing
    return None

def _gather_answers_for_category(cat: str, facility_filter: str | None = None) -> list[str]:
    """
    Collect unique answers from rows tagged with the category:
      - Accept BOTH 'g-<cat>' and 'q-<cat>' as category tags (compat with your existing data).
      - If facility_filter is provided, require an 'f-<facility>' tag that fuzzily matches it.
    """
    want = {f"g-{cat.strip().lower()}", f"q-{cat.strip().lower()}"}  # support both prefixes
    f_need = _canon(facility_filter) if facility_filter else None

    answers = []

    def matches_facility(tags: list[str]) -> bool:
        if not f_need:
            return True
        for t in tags:
            if t.startswith("f-"):
                if _canon(t[2:]) == f_need:
                    return True
        return False

    def scan(rows: list[dict]):
        for r in rows:
            tags = _tags_normalized(r.get("tags", ""))
            if want.intersection(tags) and matches_facility(tags):
                a = (r.get("answer") or "").strip()
                if a:
                    answers.append(a)

    scan(_rows_from_qa_csv())
    scan(_rows_from_client_info())

    # de-dupe, preserve order
    seen, uniq = set(), []
    for a in answers:
        if a not in seen:
            seen.add(a)
            uniq.append(a)
    return uniq

def _render_bulleted(title: str, items: list[str]) -> str:
    if not items:
        return f"No {title.lower()} found."
    return "\n".join([f"**{title}**"] + [f"- {x}" for x in items])

# ---------------- NEW: [S7.3c] “facilities with <value> as <category>” --------
def _canon_val(s: str) -> str:
    """Loose normalize for value matching (e.g., 'PCC', 'P.C.C.')."""
    if s is None: return ""
    t = str(s).lower().strip()
    t = t.translate(str.maketrans("", "", string.punctuation))
    t = re.sub(r"\s+", " ", t)
    return t

_LIST_FAC_WITH_VAL_AS_CAT = re.compile(
    r"\b(list|show)\s+facilit(?:y|ies)\s+(?:with|using)\s+(.+?)\s+(?:as|for)\s+([A-Za-z0-9 _\-]+)",
    re.IGNORECASE
)
_LIST_FAC_WHERE_CAT_IS_VAL = re.compile(
    r"\b(list|show)\s+facilit(?:y|ies)\s+where\s+([A-Za-z0-9 _\-]+)\s+(?:is|=)\s+(.+?)$",
    re.IGNORECASE
)

def _parse_facilities_by_category_value(q: str) -> tuple[str, str] | None:
    """
    Returns (category, value) if the question matches a 'facilities with X as Y' style pattern.
    Category is lowercased/singularized lightly (emr -> emr), value kept raw (e.g., 'PCC').
    """
    if not q: return None
    m = _LIST_FAC_WITH_VAL_AS_CAT.search(q)
    if m:
        value = (m.group(2) or "").strip()
        cat   = (m.group(3) or "").strip()
    else:
        m = _LIST_FAC_WHERE_CAT_IS_VAL.search(q)
        if not m:
            return None
        cat   = (m.group(2) or "").strip()
        value = (m.group(3) or "").strip()
    if len(cat) > 3 and cat.lower().endswith("s"):
        cat = cat[:-1].strip()
    return (cat.lower(), value)

def _gather_facilities_for_category_value(cat: str, value: str) -> list[str]:
    """
    Rows must have category tag: g-<cat> (or q-<cat>) AND answer matching <value>.
    Return the facilities (from f-<name> tags on those rows).
    """
    want_cat = {f"g-{cat.strip().lower()}", f"q-{cat.strip().lower()}"}
    want_val = _canon_val(value)

    facs: list[str] = []

    def scan(rows: list[dict]):
        for r in rows:
            tags = _tags_normalized(r.get("tags", ""))
            if not want_cat.intersection(tags):
                continue
            ans = (r.get("answer") or "").strip()
            if ans and _canon_val(ans) == want_val:
                for t in tags:
                    if t.startswith("f-"):
                        name = t[2:].strip()
                        if name:
                            facs.append(name)

    scan(_rows_from_qa_csv())
    scan(_rows_from_client_info())

    # de-dupe, preserve order
    seen, uniq = set(), []
    for n in facs:
        if n not in seen:
            seen.add(n)
            uniq.append(n)
    return uniq
# ------------------------------------------------------------------------------

def build_list_group_answer(q: str) -> tuple[str, list[str]] | None:
    """
    Router for List/Group mode:
      0) 'list facilities with <value> as <category>' -> facilities via g-<category> & answer==<value>
      1) 'list all facilities' -> unique facilities (f- tags + client sheet 'facility_name')
      2) 'list all <category>' -> answers from rows tagged g-<category> (or q-<category>)
         Optionally scoped to a facility if the question mentions one: "... for Aspire"
    """
    qn = (q or "").strip()

    # 0) facilities with <value> as <category>
    fv = _parse_facilities_by_category_value(qn)
    if fv:
        cat, val = fv
        facs = _gather_facilities_for_category_value(cat, val)
        if facs:
            title = f"Facilities — {cat.upper()} = {val}"
            return _render_bulleted(title, facs), []

    # 1) Facilities
    if re.search(r"\bfacilit", qn, re.IGNORECASE):
        facs = _list_unique_facilities_from_sources()
        if facs:
            return _render_bulleted("Facilities", facs), []

    # 2) Generic category (g-<category> and q-<category>)
    cat = _infer_category_from_question(qn)
    if cat:
        fac_hint = _infer_facility_from_question(qn)  # None if not mentioned
        items = _gather_answers_for_category(cat, facility_filter=fac_hint)
        if items:
            # Title: pluralize lightly
            title = cat[0].upper() + cat[1:]
            if not title.endswith("s"):
                title += "s"
            if fac_hint:
                title += f" — {fac_hint}"
            return _render_bulleted(title, items), []

    return None
















# ==============================================================================
# [S9] USER Q&A LOG — CSV + Admin routes
# ==============================================================================
# ======= User Q&A Log (CSV + helpers) =======
LOG_CSV_PATH = os.path.join("data", "user_qa_log.csv")


# [FN:_ensure_log_csv] ----------------------------------------------------
def _ensure_log_csv():
    os.makedirs(os.path.dirname(LOG_CSV_PATH), exist_ok=True)
    if not os.path.exists(LOG_CSV_PATH):
        with open(LOG_CSV_PATH, "w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow(["id","ts","section","q","a"])  # id=int, ts=ISO8601


# [FN:_read_log_rows] ----------------------------------------------------
def _read_log_rows():
    _ensure_log_csv()
    rows = []
    with open(LOG_CSV_PATH, "r", newline="", encoding="utf-8") as f:
        r = csv.DictReader(f)
        for row in r:
            rows.append({
                "id": int(row.get("id") or 0),
                "ts": (row.get("ts") or "").strip(),
                "section": (row.get("section") or "").strip(),
                "q": (row.get("q") or "").strip(),
                "a": (row.get("a") or "").strip(),
            })
    return rows


# [FN:_next_log_id] ----------------------------------------------------
def _next_log_id(rows):
    return (max([x["id"] for x in rows] or [0]) + 1)


# [FN:_append_user_qa_log] ----------------------------------------------------
def _append_user_qa_log(section: str, q: str, a: str):
    _ensure_log_csv()
    rows = _read_log_rows()
    new_id = _next_log_id(rows)
    from datetime import datetime, timezone
    ts = datetime.now(timezone.utc).isoformat()
    # decide id first
    if USE_SQLITE_WRITE:
        new_id = _next_log_id_sqlite()
    else:
        rows = _read_log_rows()
        new_id = _next_log_id(rows)

    ts = datetime.now(timezone.utc).isoformat()

    if USE_SQLITE_WRITE:
        # primary write → SQLite
        upsert_log_sqlite({
            "id": new_id, "ts": ts, "section": (section or "").strip(),
            "q": (q or "").strip(), "a": (a or "").strip(),
            "promoted": 0, "a_quality": ""
        })
        # optional CSV mirror
        if MIRROR_SQLITE_WRITES:
            with open(LOG_CSV_PATH, "a", newline="", encoding="utf-8") as f:
                w = csv.writer(f)
                w.writerow([new_id, ts, (section or "").strip(), (q or "").strip(), (a or "").strip(), 0, ""])
    else:
        # primary write → CSV
        with open(LOG_CSV_PATH, "a", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow([new_id, ts, (section or "").strip(), (q or "").strip(), (a or "").strip(), 0, ""])
        # mirror to SQLite
        if MIRROR_SQLITE_WRITES:
            try:
                upsert_log_sqlite({
                    "id": new_id, "ts": ts, "section": (section or "").strip(),
                    "q": (q or "").strip(), "a": (a or "").strip(),
                    "promoted": 0, "a_quality": ""
                })
            except Exception as e:
                print("mirror user_qa_log add->sqlite failed:", e)

    return new_id


# --- Extend log schema with 'promoted' flag (backward-compatible) ---

def _write_log_rows(rows: list[dict]):
    _ensure_log_csv()
    with open(LOG_CSV_PATH, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["id","ts","section","q","a","promoted"])
        for r in rows:
            w.writerow([
                int(r.get("id", 0)),
                (r.get("ts") or "").strip(),
                (r.get("section") or "").strip(),
                (r.get("q") or "").strip(),
                (r.get("a") or "").strip(),
                1 if str(r.get("promoted", "")).strip() in ("1","true","True") else 0
            ])

def _read_log_rows():
    _ensure_log_csv()
    rows = []
    with open(LOG_CSV_PATH, "r", newline="", encoding="utf-8") as f:
        r = csv.DictReader(f)
        for row in r:
            rows.append({
                "id": int(row.get("id") or 0),
                "ts": (row.get("ts") or "").strip(),
                "section": (row.get("section") or "").strip(),
                "q": (row.get("q") or "").strip(),
                "a": (row.get("a") or "").strip(),
                "promoted": 1 if str(row.get("promoted","")).strip() in ("1","true","True") else 0
            })
    return rows

def _ensure_log_csv():
    os.makedirs(os.path.dirname(LOG_CSV_PATH), exist_ok=True)
    if not os.path.exists(LOG_CSV_PATH):
        with open(LOG_CSV_PATH, "w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow(["id","ts","section","q","a","promoted"])
    else:
        # upgrade file missing 'promoted' column
        with open(LOG_CSV_PATH, "r", newline="", encoding="utf-8") as f:
            first = f.readline()
        if "promoted" not in first:
            rows = []
            with open(LOG_CSV_PATH, "r", newline="", encoding="utf-8") as f:
                r = csv.DictReader(f)
                for row in r:
                    rows.append({
                        "id": int(row.get("id") or 0),
                        "ts": (row.get("ts") or "").strip(),
                        "section": (row.get("section") or "").strip(),
                        "q": (row.get("q") or "").strip(),
                        "a": (row.get("a") or "").strip(),
                        "promoted": 0
                    })
            _write_log_rows(rows)

def _mark_log_promoted(log_id: int, value: bool = True):
    rows = _read_log_rows()
    for r in rows:
        if int(r["id"]) == int(log_id):
            r["promoted"] = 1 if value else 0
            break
    _write_log_rows(rows)

TOKEN_RE = re.compile(r"\[([^\]]+)\]")


# [FN:_best_facility_name] ----------------------------------------------------
def _best_facility_name(row_dict: dict) -> str:
    for k in ("legal_name", "facility_name", "doing_business_as", "dba"):
        if k in row_dict and str(row_dict[k]).strip():
            return str(row_dict[k]).strip()
    return ""


# [FN:_render_question] ----------------------------------------------------
def _render_question(template: str, row_dict: dict, header_map: dict) -> str:
    """Replace [Tokens] in Row-2 template using the submission row values."""
    def repl(m):
        token = m.group(1).strip()
        if token.lower() in ("facility", "facility name", "client"):
            val = _best_facility_name(row_dict)
            return val if val else m.group(0)
        key = header_map.get(token.lower())
        if key and str(row_dict.get(key, "")).strip():
            return str(row_dict[key]).strip()
        return m.group(0)
    return TOKEN_RE.sub(repl, str(template or ""))


# [FN:_make_qid] ----------------------------------------------------
def _make_qid(facility: str, field_name: str) -> str:
    return f"{slugify(facility) or 'client'}:{field_name}"


# [FN:_embed_texts] ----------------------------------------------------
def _embed_texts(chunks: list[str]) -> list[list[float]]:
    if not chunks:
        return []
    resp = client.embeddings.create(model=EMBED_MODEL, input=chunks)
    return [d.embedding for d in resp.data]


# [FN:_chroma] ----------------------------------------------------
def _chroma() -> chromadb.Collection:
    os.makedirs(VECTOR_PATH, exist_ok=True)
    chroma = chromadb.PersistentClient(path=VECTOR_PATH)
    return chroma.get_or_create_collection(name=QA_COLLECTION)


# [FN:_upsert_into_chroma] ----------------------------------------------------
def _upsert_into_chroma(rows: list[dict]):
    """
    rows: [{id, question, answer, ...}] -> embeds (question || answer) and upserts IDs only for these rows.
    NOTE: This matches your build_index.py doc pattern (question + ' || ' + answer). 
    """
    if not rows:
        return
    ids = [r["id"] for r in rows]
    docs = [(r["question"] + " || " + r["answer"]) for r in rows]  # same as build_index.py docs format
    metas = [{"question": r["question"], "answer": r["answer"]} for r in rows]

    embs = _embed_texts(docs)
    coll = _chroma()

    # Chroma upsert: safest across versions is delete+add for these IDs
    try:
        coll.delete(ids=ids)
    except Exception:
        pass
    coll.add(ids=ids, documents=docs, embeddings=embs, metadatas=metas)


# [FN:_upsert_into_qa_csv] ----------------------------------------------------
def _upsert_into_qa_csv(new_rows: list[dict]):
    """
    Upsert by 'id' into data/qa.csv (replace any existing IDs with new ones).
    new_rows items must have keys: id, section, question, answer, tags
    """
    if not new_rows:
        return

    try:
        qa = pd.read_csv(QA_CSV_PATH).fillna("")
    except FileNotFoundError:
        qa = pd.DataFrame(columns=["id", "section", "question", "answer", "tags"])

    new_df = pd.DataFrame(new_rows, columns=["id","section","question","answer","tags"])
    if not qa.empty:
        qa = qa[~qa["id"].isin(new_df["id"])]
    qa = pd.concat([new_df, qa], ignore_index=True)
    qa.to_csv(QA_CSV_PATH, index=False)


# [FN:compute the next ID] ----------------------------------------------------
# === Auto-ID helpers & lock (server is the source of truth) ===
QA_LOCK = threading.Lock()

def _parse_id_number(id_str: str) -> int:
    """Extract numeric digits from an ID like 'Q007' -> 7, 'Q20251107-001' -> 20251107001."""
    s = str(id_str or "")
    nums = re.findall(r"\d+", s)
    return int("".join(nums)) if nums else 0

def _next_id_from_rows(rows, prefix="Q", width=3) -> str:
    """
    Sequential classic style: Q001, Q002, ...
    Finds max numeric and returns next, zero-padded.
    """
    mx = 0
    for r in rows or []:
        mx = max(mx, _parse_id_number(r.get("id")))
    return f"{prefix}{str(mx + 1).zfill(width)}"



# [FN:_is_blank] ----------------------------------------------------
def _is_blank(v) -> bool:
    if v is None:
        return True
    s = str(v).strip()
    return s == "" or s.lower() in {"nan", "none", "null", "na", "n/a", "-" , "--"}


# [FN:_clean_str] ----------------------------------------------------
def _clean_str(v) -> str:
    return "" if _is_blank(v) else str(v).strip()


# [FN:_generate_qas_for_submission] ----------------------------------------------------
def _generate_qas_for_submission(row1: list[str], row2: list[str], submission: pd.Series) -> list[dict]:
    """
    Build Q&A rows ONLY for this single submission (Row 3+), using row1 field names and row2 question templates.
    Returns: list of dicts with {id, section, question, answer, tags}
    """

    # NEW: Load optional Row 3 per-column tags from the sheet on disk
    col_tags = []
    try:
        df_head = pd.read_csv(CLIENT_SHEET_PATH, header=None, nrows=3).fillna("")
        if df_head.shape[0] >= 3:
            col_tags = df_head.iloc[2].astype(str).tolist()  # Row 3 tags
    except Exception:
        col_tags = []

    # Build a dict of field_name -> cleaned value (no 'nan', 'None', etc.)
    row_dict = {
        row1[i]: _clean_str(submission.iloc[i]) if i < len(submission) else ""
        for i in range(len(row1))
    }


    # Column B = section, Column C = tags (cleaned)
    section = _clean_str(submission.iloc[1]) if len(submission) > 1 else "Facility Details"
    tags    = _clean_str(submission.iloc[2]) if len(submission) > 2 else ""

    # case-insensitive map for tokens: [City] -> 'city'
    header_map = {}
    for name in row1:
        if name:
            header_map[name.strip().lower()] = name

    facility = _best_facility_name(row_dict) or "Client"
    out = []
    # === ID seed so we can emit Q001-style IDs like /admin/add ===
    with QA_LOCK:
        _df_existing = load_csv_df()
        _rows_existing = _df_existing.to_dict(orient="records")
        _seed = 0
        for _r in _rows_existing or []:
            _seed = max(_seed, _parse_id_number(_r.get("id")))

    START_COL = 3  # 0=A, 1=B (Section), 2=C (Tags), 3=D (first real question)
    for col_idx, template in enumerate(row2[START_COL:], start=START_COL):
        if not template or col_idx >= len(row1):
            continue

        field_name = row1[col_idx]
        # Only include if the field was actually provided/edited (i.e., not blank/NaN)
        ans_raw = submission.iloc[col_idx] if col_idx < len(submission) else ""
        if _is_blank(ans_raw):
            continue

        ans = _clean_str(ans_raw)
        q   = _render_question(template, row_dict, header_map)
        
        _seed += 1
        qid = f"Q{str(_seed).zfill(3)}"

        # Keep Column C (global tags) + optional Row 3 per-column tag
        tags_this = tags
        if col_idx < len(col_tags) and str(col_tags[col_idx]).strip():
            t_extra = str(col_tags[col_idx]).strip()
            tags_this = ", ".join([t for t in [tags, t_extra] if t]).strip(", ").strip()

        out.append({
            "id": qid,
            "section": section or "Facility Details",
            "question": q,
            "answer": ans,
            "tags": tags_this,
        })



    return out


# [FN:_build_tags_from_payload] ----------------------------------------------------
def _build_tags_from_payload(payload: dict) -> str:
    # Use client-provided hint if present
    h = str(payload.get("__tag_hint", "")).strip()
    if h:
        return h
    name = (payload.get("legal_name") or payload.get("facility_name") or "").strip()
    parts = []
    if name:
        parts.append(f"f-{name}")
    return ", ".join([p for p in parts if p])


# ---- Facilities Admin routes --------------------------------------------------

@app.post("/admin/facilities/create")
def admin_facilities_create(
    request: Request,
    facility_name: str = Form(...),
    facility_id: str = Form(...),
    corporate_group: str = Form("")
):
    """
    Create a new facility master record and return its intake URL.
    Requires admin token via header `x-admin-token` or ?token=...
    """
    require_admin(request)
    from datetime import datetime, timezone
    import sqlite3 as _sqlite3

    conn = _db()
    cur = conn.cursor()

    # 1) Basic validation
    name_clean  = (facility_name or "").strip()
    id_input    = (facility_id or "").strip()
    group_clean = (corporate_group or "").strip()
    if not name_clean or not id_input:
        raise HTTPException(status_code=400, detail="facility_name and facility_id are required.")

    # 2) Enforce slug format: lowercase letters, digits, hyphens
    id_slug = slugify(id_input) if 'slugify' in globals() else re.sub(r"[^a-z0-9]+", "-", id_input.lower()).strip("-")
    if id_slug != id_input:
        # Tell the user exactly what to use
        raise HTTPException(
            status_code=400,
            detail=f"facility_id must use lowercase letters, numbers, and hyphens. Suggested: '{id_slug}'."
        )

    # 3) Pre-check uniqueness (nice error if the ID already exists)
    cur.execute("SELECT 1 FROM facilities WHERE facility_id = ? LIMIT 1;", (id_slug,))
    if cur.fetchone():
        raise HTTPException(status_code=409, detail="Facility ID already exists. Choose a different ID.")

    # 4) Create token + timestamps
    token = _gen_token(24)
    now   = datetime.now(timezone.utc).isoformat()

    # 5) Insert row with hard guard: catch UNIQUE collisions and return 409 (not 500)
    try:
        cur.execute("""
            INSERT INTO facilities (
                facility_id, facility_name, corporate_group,
                intake_token, intake_status, created_at, updated_at, extras
            )
            VALUES (?,?,?,?, 'not-started', ?, ?, NULL);
        """, (id_slug, name_clean, group_clean, token, now, now))
        conn.commit()
    except _sqlite3.IntegrityError:
        # In case of a race (two admins submit same ID), return a friendly conflict
        raise HTTPException(status_code=409, detail="Facility ID already exists. Choose a different ID.")

    # 6) Build URLs to return — always use the cleaned/validated slug
    urls = _facility_intake_urls(request, id_slug, token)

    return {
        "ok": True,
        "facility": {
            "facility_id":   id_slug,
            "facility_name": name_clean,
            "corporate_group": group_clean,
            "intake_status": "not-started",
            "intake_token":  token,
            "created_at":    now,
            "updated_at":    now,
            "url_full":      urls["full"],
            "url_path":      urls["path"],
        }
    }



@app.get("/admin/facilities")
def admin_facilities_list(request: Request, q: str = Query(default="")):
    """
    List facilities (optionally filter by name/corporate group).
    Requires admin token via header `x-admin-token` or ?token=...
    """
    require_admin(request)
    conn = _db()
    cur = conn.cursor()
    q_like = f"%{(q or '').strip().lower()}%"
    cur.execute("""
        SELECT facility_id, facility_name, corporate_group, intake_token, intake_status, created_at, updated_at
        FROM facilities
        WHERE (LOWER(facility_name) LIKE ? OR LOWER(IFNULL(corporate_group,'')) LIKE ?)
        ORDER BY LOWER(facility_name) ASC;
    """, (q_like, q_like))
    rows = cur.fetchall()

    out = []
    for (fid, name, group, tok, status, ca, ua) in rows:
        urls = _facility_intake_urls(request, fid, tok or "")
        out.append({
            "facility_id": fid,
            "facility_name": name,
            "corporate_group": group or "",
            "intake_status": status or "not-started",
            "created_at": ca,
            "updated_at": ua,
            "url_full": urls["full"] if tok else "",
            "url_path": urls["path"] if tok else "",
            "intake_token": tok or "",
            "wix_url": (f"https://www.startevolv.com/facility-details?f={fid}&t={tok}") if tok else "",
        })
    return {"ok": True, "items": out}

@app.get("/admin/facilities/{facility_id}")
def admin_facility_detail(request: Request, facility_id: str):
    """
    Return a single facility row for admin editing.
    Requires admin token via header `x-admin-token` or ?token=...
    """
    require_admin(request)
    conn = _db()
    cur = conn.cursor()
    cur.execute("""
        SELECT facility_id, facility_name, corporate_group, intake_token, intake_status, created_at, updated_at, extras
        FROM facilities
        WHERE facility_id = ? LIMIT 1;
    """, (facility_id,))
    row = cur.fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Facility not found")

    fid, name, group, tok, status, ca, ua, extras = row
    urls = _facility_intake_urls(request, fid, tok or "")
    return {
        "ok": True,
        "facility": {
            "facility_id": fid,
            "facility_name": name,
            "corporate_group": group or "",
            "intake_status": status or "not-started",
            "created_at": ca,
            "updated_at": ua,
            "extras": extras,
            "url_full": urls["full"] if tok else "",
            "url_path": urls["path"] if tok else "",
            "intake_token": tok or "",
            "wix_url": (f"https://www.startevolv.com/facility-details?f={fid}&t={tok}") if tok else "",
        }
    }


@app.post("/admin/facilities/update")
def admin_facility_update(
    request: Request,
    facility_id: str = Form(...),
    facility_name: str = Form(""),
    corporate_group: str = Form(""),
    intake_status: str = Form("")
):
    """
    Update basic facility fields (name, group, status).
    Requires admin token via header `x-admin-token` or ?token=...
    """
    require_admin(request)
    from datetime import datetime, timezone
    now = datetime.now(timezone.utc).isoformat()

    conn = _db()
    cur = conn.cursor()

    # Ensure the row exists
    cur.execute("SELECT 1 FROM facilities WHERE facility_id = ? LIMIT 1;", (facility_id,))
    if not cur.fetchone():
        raise HTTPException(status_code=404, detail="Facility not found")

    # Build dynamic update set
    fields = []
    params = []
    if facility_name:
        fields.append("facility_name = ?")
        params.append(facility_name.strip())
    if corporate_group:
        fields.append("corporate_group = ?")
        params.append(corporate_group.strip())
    if intake_status:
        fields.append("intake_status = ?")
        params.append(intake_status.strip())

    # Always bump updated_at
    fields.append("updated_at = ?")
    params.append(now)

    params.append(facility_id)
    cur.execute(f"UPDATE facilities SET {', '.join(fields)} WHERE facility_id = ?;", params)
    conn.commit()

    return {"ok": True, "updated_at": now}

# ===== Quick Facts (Admin) =====
from fastapi import Form, Query, Request

@app.get("/admin/facilities/{fid}/facts")
def admin_facility_facts_list(request: Request, fid: str):
    """
    List quick facts for a facility.
    Token accepted via header x-admin-token or ?token=...  (handled by require_admin)
    """
    require_admin(request)
    con = _db()
    cur = con.execute(
        "SELECT id, facility_id, fact_text, tags, created_at "
        "FROM fac_facts WHERE facility_id=? ORDER BY id DESC;",
        (fid,)
    )
    cols = [c[0] for c in cur.description]
    rows = [dict(zip(cols, r)) for r in cur.fetchall()]
    con.close()
    # Align with Admin HTML: it expects each item to contain fact_text, tags, created_at
    return {"ok": True, "items": rows}

@app.post("/admin/facilities/facts/add")
def admin_facility_facts_add(
    request: Request,
    facility_id: str = Form(...),
    fact_text: str  = Form(...),
    tags: str       = Form("")
):
    """
    Add a quick fact row.
    """
    require_admin(request)
    txt = (fact_text or "").strip()
    if not txt:
        raise HTTPException(status_code=400, detail="Fact text is required.")
    con = _db()
    cur = con.execute(
        "INSERT INTO fac_facts (facility_id, fact_text, tags) VALUES (?, ?, ?);",
        (facility_id, txt, (tags or "").strip())
    )
    new_id = cur.lastrowid
    con.commit()
    row = con.execute(
        "SELECT id, facility_id, fact_text, tags, created_at FROM fac_facts WHERE id=?;",
        (new_id,)
    ).fetchone()
    con.close()
    return {"ok": True, "item": {
        "id": row[0], "facility_id": row[1], "fact_text": row[2], "tags": row[3], "created_at": row[4]
    }}

@app.post("/admin/facilities/facts/delete")
def admin_facility_facts_delete(
    request: Request,
    id: int = Form(...)
):
    """
    Delete a quick fact row by id.
    """
    require_admin(request)
    con = _db()
    con.execute("DELETE FROM fac_facts WHERE id=?;", (int(id),))
    con.commit()
    con.close()
    return {"ok": True}




# ---- User Q&A Log Admin routes ----
from datetime import datetime

@app.get("/admin/ulog/list")

# [FN:admin_ulog_list] ---------------------------------------------------
def admin_ulog_list(
    request: Request,
    from_: str = Query(default="", alias="from"),
    to: str = Query(default=""),
    section: str = Query(default=""),
    topics: str = Query(default=""),       # NEW: topics filter (string)
    q: str = Query(default=""),
    quality: str = Query(default="")       # "", "up", or "down"
):

    require_admin(request)
    rows = _read_log_rows()  # [{id, ts, section, q, a}]

    def parse_date(s):
        try:
            return datetime.fromisoformat(s)
        except Exception:
            return None

    # filters
    out = rows
    if from_:
        df = parse_date(from_ + "T00:00:00")
        if df:
            out = [r for r in out if parse_date((r["ts"] or "").split("Z")[0]) and parse_date((r["ts"] or "").split("Z")[0]) >= df]
    if to:
        dt = parse_date(to + "T23:59:59")
        if dt:
            out = [r for r in out if parse_date((r["ts"] or "").split("Z")[0]) and parse_date((r["ts"] or "").split("Z")[0]) <= dt]
    if section:
        out = [r for r in out if (r.get("section") or "").strip() == section.strip()]

    # NEW: topics param behaves as a "contains" filter on the stored topics string (in 'section')
    if topics:
        tpc = topics.strip().lower()
        out = [r for r in out if tpc and tpc in (r.get("section", "") or "").lower()]

    if q:
        ql = q.lower().strip()
        out = [r for r in out if ql in (r.get("q","").lower()+" "+r.get("a","").lower())]
    
    # quality filter
    if quality:
        qv = quality.strip().lower()
        out = [r for r in out if (r.get("a_quality","").strip().lower() == qv)]
    
    # Sort newest first (id is monotonic in writer)
    out.sort(key=lambda r: r.get("id", 0), reverse=True)

    # Provide 'topics' alias for convenience (mirrors the stored 'section' string)
    for r in out:
        if "topics" not in r:
            r["topics"] = r.get("section", "")

    return {"rows": out}


@app.post("/admin/ulog/delete")

@app.get("/admin/ulog/get")
def admin_ulog_get(request: Request, id: int = Query(...)):
    require_admin(request)
    rows = _read_log_rows()
    for r in rows:
        if int(r["id"]) == int(id):
            # add 'topics' alias mirroring 'section'
            if "topics" not in r:
                r["topics"] = r.get("section", "")
            return {"row": r}
    raise HTTPException(status_code=404, detail="Log row not found")


@app.post("/admin/ulog/promote")
def admin_ulog_promote(
    request: Request,
    id: int = Form(...),                 # log row id to promote
    new_id: str = Form(""),              # optional override for qa.csv id
    section: str = Form(""),             # legacy field; we still accept it
    topics: str = Form(""),              # NEW: topics from the Promote modal
    question: str = Form(""),            # edited question (fallback to log.q)
    answer: str = Form(""),              # edited answer   (fallback to log.a)
    tags: str = Form("")                 # optional tags
):

    require_admin(request)

    # 1) find the log row
    rows = _read_log_rows()
    log_row = next((r for r in rows if int(r["id"]) == int(id)), None)
    if not log_row:
        raise HTTPException(status_code=404, detail="Log row not found")

    # 2) materialize final fields (edited values win, then log fallback)
    #    NOTE: we treat 'topics' as the new field; it maps to the stored 'section' column.
    section_f = (topics or section or log_row.get("section") or "").strip()
    question_f = (question or log_row.get("q") or "").strip()
    answer_f = (answer or log_row.get("a") or "").strip()
    tags_f = (tags or "").strip()


    if not question_f or not answer_f:
        raise HTTPException(status_code=400, detail="Question and Answer are required")

    # 3) choose id (generate if missing or collides)
    df = load_qa_df_unified()
    existing = set(df["id"].astype(str).tolist())
    qid = (new_id or "").strip()
    if not qid or qid in existing:
        qid = _next_id(existing, prefix="Q")  # [FN:_next_id] lives under [S16]

    # 4) write new Q&A (primary → SQLite when USE_SQLITE_WRITE=1)
    new_row = {
        "id": qid,
        "section": section_f,
        "question": question_f,
        "answer": answer_f,
        "tags": tags_f,
    }

    # Prepare SQLite shape (qa table has a 'topics' column)
    new_row_sql = dict(new_row)
    new_row_sql["topics"] = new_row_sql.get("section", "")

    if USE_SQLITE_WRITE:
        # Primary write → SQLite
        upsert_qa_sqlite(new_row_sql)
        # Optional CSV mirror
        if MIRROR_SQLITE_WRITES:
            df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
            save_csv_df(df)
    else:
        # Primary write → CSV (legacy)
        df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
        save_csv_df(df)
        # Mirror to SQLite
        if MIRROR_SQLITE_WRITES:
            try:
                upsert_qa_sqlite(new_row_sql)
            except Exception as e:
                print("mirror promote qa->sqlite failed:", e)

    # 5) rebuild index so it’s immediately searchable (CSV indexer for now)
    try:
        rebuild_index_from_csv()
    except Exception as e:
        print("Reindex after promote failed:", e)

    # 6) mark the log row as promoted
    if USE_SQLITE_WRITE:
        # Primary → SQLite log
        upsert_log_sqlite({
            "id": int(id),
            "ts": log_row.get("ts",""),
            "section": log_row.get("section",""),
            "q": log_row.get("q",""),
            "a": log_row.get("a",""),
            "promoted": 1,
            "a_quality": log_row.get("a_quality",""),
        })
        # Optional CSV mirror
        if MIRROR_SQLITE_WRITES:
            _mark_log_promoted(id, True)
    else:
        # Primary → CSV log
        _mark_log_promoted(id, True)
        # Mirror to SQLite
        if MIRROR_SQLITE_WRITES:
            try:
                upsert_log_sqlite({
                    "id": int(id),
                    "ts": log_row.get("ts",""),
                    "section": log_row.get("section",""),
                    "q": log_row.get("q",""),
                    "a": log_row.get("a",""),
                    "promoted": 1,
                    "a_quality": log_row.get("a_quality",""),
                })
            except Exception as e:
                print("mirror log promote->sqlite failed:", e)

    return {"ok": True, "added": new_row}






# [FN:admin_ulog_delete] ----------------------------------------------------
def admin_ulog_delete(request: Request, id: int = Form(...)):
    require_admin(request)
    rows = _read_log_rows()
    new_rows = [r for r in rows if int(r.get("id",0)) != int(id)]
    # rewrite file (now with promoted column)
    _write_log_rows(new_rows)
    return {"ok": True, "remaining": len(new_rows)}



# --- Set quality (public; no admin token) ---
@app.post("/ulog/quality")
def ulog_quality_public(id: int = Form(...), quality: str = Form(...)):
    """
    Public endpoint used by the Copilot ask page.
    quality: "up" | "down" | "" (empty clears)
    """
    rows = _read_log_rows()
    found = False
    qv = quality.strip().lower()
    if qv not in ("up","down",""):
        raise HTTPException(status_code=400, detail="quality must be 'up', 'down', or ''")
    for r in rows:
        if int(r["id"]) == int(id):
            r["a_quality"] = qv
            found = True
            break
    if not found:
        raise HTTPException(status_code=404, detail="Log row not found")
    _write_log_rows(rows)
    # NEW (Step 3.3): mirror quality change to SQLite
    if MIRROR_SQLITE_WRITES:
        try:
            # find the updated row
            row = next((r for r in rows if int(r["id"]) == int(id)), None)
            if row:
                upsert_log_sqlite(row)
        except Exception as e:
            print("mirror log quality->sqlite failed:", e)
    return {"ok": True, "id": id, "a_quality": qv}

# --- Set quality (admin; mirrors public but requires token) ---
@app.post("/admin/ulog/quality")
def ulog_quality_admin(request: Request, id: int = Form(...), quality: str = Form(...)):
    require_admin(request)
    return ulog_quality_public(id=id, quality=quality)





# =============================
# [PATCH S9] Fix log CSV helpers to avoid recursion
# =============================
# Root cause: earlier versions called _ensure_log_csv() -> _write_log_rows() which
# in turn called _ensure_log_csv() again during "upgrade" to add the 'promoted' column,
# causing infinite recursion. These re-definitions replace the helpers with a
# non-recursive, idempotent upgrade path. (Last definition wins in Python.)

LOG_HEADERS = ["id","ts","section","q","a","promoted","a_quality"]

def _ensure_log_csv():
    import os, csv
    os.makedirs(os.path.dirname(LOG_CSV_PATH), exist_ok=True)
    # Create new file with correct headers
    if not os.path.exists(LOG_CSV_PATH):
        with open(LOG_CSV_PATH, "w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow(LOG_HEADERS)
        return

    # Upgrade if file exists but is missing promoted or a_quality columns
    with open(LOG_CSV_PATH, "r", newline="", encoding="utf-8") as f:
        first_line = f.readline()

    need_upgrade = ("promoted" not in first_line) or ("a_quality" not in first_line)
    if need_upgrade:
        rows = []
        with open(LOG_CSV_PATH, "r", newline="", encoding="utf-8") as f:
            r = csv.DictReader(f)
            for row in r:
                rows.append({
                    "id": int(row.get("id") or 0),
                    "ts": (row.get("ts") or "").strip(),
                    "section": (row.get("section") or "").strip(),
                    "q": (row.get("q") or "").strip(),
                    "a": (row.get("a") or "").strip(),
                    "promoted": 1 if str(row.get("promoted","")).strip() in ("1","true","True") else 0,
                    "a_quality": (row.get("a_quality") or "").strip(),  # keep if present, else empty
                })
        with open(LOG_CSV_PATH, "w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow(LOG_HEADERS)
            for r in rows:
                w.writerow([
                    int(r.get("id", 0)),
                    (r.get("ts") or "").strip(),
                    (r.get("section") or "").strip(),
                    (r.get("q") or "").strip(),
                    (r.get("a") or "").strip(),
                    1 if str(r.get("promoted", "")).strip() in ("1","true","True") else 0,
                    (r.get("a_quality") or "").strip(),
                ])

def _read_log_rows():
    import csv
    _ensure_log_csv()
    rows = []
    with open(LOG_CSV_PATH, "r", newline="", encoding="utf-8") as f:
        r = csv.DictReader(f)
        for row in r:
            rows.append({
                "id": int(row.get("id") or 0),
                "ts": (row.get("ts") or "").strip(),
                "section": (row.get("section") or "").strip(),
                "q": (row.get("q") or "").strip(),
                "a": (row.get("a") or "").strip(),
                "promoted": 1 if str(row.get("promoted","")).strip() in ("1","true","True") else 0,
                "a_quality": (row.get("a_quality") or "").strip(),  # "", "up", "down"
            })
    return rows

def _write_log_rows(rows: list[dict]):
    import os, csv
    os.makedirs(os.path.dirname(LOG_CSV_PATH), exist_ok=True)
    with open(LOG_CSV_PATH, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(LOG_HEADERS)
        for r in rows:
            w.writerow([
                int(r.get("id", 0)),
                (r.get("ts") or "").strip(),
                (r.get("section") or "").strip(),
                (r.get("q") or "").strip(),
                (r.get("a") or "").strip(),
                1 if str(r.get("promoted", "")).strip() in ("1","true","True") else 0,
                (r.get("a_quality") or "").strip(),
            ])

def _append_user_qa_log(section: str, q: str, a: str):
    import csv
    from datetime import datetime, timezone
    _ensure_log_csv()
    rows = _read_log_rows()
    new_id = _next_log_id(rows)
    ts = datetime.now(timezone.utc).isoformat()
    with open(LOG_CSV_PATH, "a", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        #                                promoted a_quality
        w.writerow([new_id, ts, (section or "").strip(), (q or "").strip(), (a or "").strip(), 0, ""])
    # NEW (Step 3.3): mirror log write to SQLite
    if MIRROR_SQLITE_WRITES:
        try:
            upsert_log_sqlite({
                "id": new_id, "ts": ts, "section": (section or "").strip(),
                "q": (q or "").strip(), "a": (a or "").strip(),
                "promoted": 0, "a_quality": ""
            })
        except Exception as e:
            print("mirror user_qa_log add->sqlite failed:", e)
    return new_id
# =============================
# [END PATCH S9]
# =============================
# ==============================================================================
# [S10] CHROMA GUARDS & SELF-HEAL HELPERS
# ==============================================================================
# ---- Chroma guards & self-heal helpers ----
from chromadb.errors import InternalError as ChromaInternalError


# [FN:_count_safe] ----------------------------------------------------
def _count_safe(coll):
    try:
        return coll.count()
    except Exception:
        return 0


# [FN:_safe_query] ----------------------------------------------------
def _safe_query(coll, vvec, n_results, include, rebuild_fn=None):
    # Skip if collection is empty
    if _count_safe(coll) == 0:
        return {"ids":[[]], "documents":[[]], "metadatas":[[]], "distances":[[]], "embeddings":[[]]}
    try:
        return coll.query(query_embeddings=[vvec], n_results=n_results, include=include)
    except ChromaInternalError:
        # Self-heal once: rebuild then retry
        if rebuild_fn:
            try:
                rebuild_fn()
            except Exception:
                pass
        try:
            return coll.query(query_embeddings=[vvec], n_results=n_results, include=include)
        except Exception:
            # Final fallback: return empty results
            return {"ids":[[]], "documents":[[]], "metadatas":[[]], "distances":[[]], "embeddings":[[]]}




# ==============================================================================
# [S10.1] REBUILD QA INDEX FROM CSV
# ==============================================================================

# [FN:rebuild_index_from_csv] ----------------------------------------------------
def rebuild_index_from_csv():
    df = load_csv_df()
    docs  = (df["question"] + " || " + df["answer"]).astype(str).tolist()
    ids   = df["id"].astype(str).tolist()
    # include topics (plus keep section for back-compat)
    metas = df[["section", "topics", "question", "answer", "tags"]].to_dict(orient="records")


    emb = client.embeddings.create(model=EMBED_MODEL, input=docs)
    embs = [d.embedding for d in emb.data]

    if collection.count() > 0:
        try:
            existing = collection.get()
            if existing and existing.get("ids"):
                collection.delete(ids=existing["ids"])
        except Exception:
            pass

    collection.add(ids=ids, documents=docs, embeddings=embs, metadatas=metas)
    return len(ids)



# ==============================================================================
# [S11] REBUILD DOCS INDEX FROM /docs
# ==============================================================================

# [FN:rebuild_docs_index_from_folder] ----------------------------------------------------
def rebuild_docs_index_from_folder():
    """
    Rebuild the docs_main collection from files in backend/docs.
    Supports .txt/.md; will also read .pdf and .docx if those parsers are installed.
    """
    DOCS_DIR = os.path.join(os.path.dirname(__file__), "docs")
    os.makedirs(DOCS_DIR, exist_ok=True)

    def chunk_text(text: str, chunk_chars=3500, overlap=400):
        text = (text or "").strip()
        if not text:
            return []
        out = []
        start = 0
        n = len(text)
        while start < n:
            end = min(n, start + chunk_chars)
            slice_ = text[start:end]
            cut = slice_.rfind("\n\n")
            if cut > 1200:
                end = start + cut
                slice_ = text[start:end]
            out.append(slice_.strip())
            if end == n:
                break
            start = max(end - overlap, 0)
        return [s for s in out if s]

    def read_txt(path):
        return open(path, "r", encoding="utf-8", errors="ignore").read()

    def read_pdf(path):
        try:
            from pypdf import PdfReader
        except Exception:
            return ""  # parser not installed
        parts = []
        try:
            reader = PdfReader(path)
            for i, page in enumerate(reader.pages, 1):
                try:
                    txt = page.extract_text() or ""
                except Exception:
                    txt = ""
                if txt.strip():
                    parts.append(f"[PDF Page {i}]\n{txt}")
        except Exception:
            return ""
        return "\n\n".join(parts)

    def read_docx(path):
        try:
            import docx
        except Exception:
            return ""  # parser not installed
        try:
            d = docx.Document(path)
            paras = [p.text for p in d.paragraphs if p.text and p.text.strip()]
            return "\n".join(paras)
        except Exception:
            return ""

    docs = []
    # gather files
    for name in os.listdir(DOCS_DIR):
        fp = os.path.join(DOCS_DIR, name)
        if not os.path.isfile(fp):
            continue
        ext = os.path.splitext(name)[1].lower()
        if ext in (".txt", ".md"):
            raw = read_txt(fp)
        elif ext == ".pdf":
            raw = read_pdf(fp)
        elif ext == ".docx":
            raw = read_docx(fp)
        else:
            continue  # unsupported type; add more as needed

        raw = raw.strip()
        if not raw:
            continue

        chunks = chunk_text(raw, chunk_chars=3500, overlap=400)
        for idx, ch in enumerate(chunks):
            docs.append({
                "id": f"{name}:{idx}",
                "text": ch,
                "meta": {"source": name, "chunk": idx}
            })

    if not docs:
        return {"ok": True, "count": 0, "note": "No parsable documents found in docs/"}

    # embed
    embs_resp = client.embeddings.create(model=EMBED_MODEL, input=[d["text"] for d in docs])
    embs = [d.embedding for d in embs_resp.data]

    # replace collection
    if docs_collection.count() > 0:
        try:
            existing = docs_collection.get()
            if existing and existing.get("ids"):
                docs_collection.delete(ids=existing["ids"])
        except Exception:
            pass

    docs_collection.add(
        ids=[d["id"] for d in docs],
        documents=[d["text"] for d in docs],
        metadatas=[d["meta"] for d in docs],
        embeddings=embs
    )
    return {"ok": True, "count": len(docs)}

    


# ==============================================================================
# [S12] HEALTH & ROOT ROUTES (simple)
# ==============================================================================
# =========================
# Routes
# =========================
@app.get("/health")

# [FN:health] ----------------------------------------------------
def health():
    return {
        "ok": True,
        "qa_count": collection.count(),
        "docs_count": docs_collection.count(),
    }
    
@app.get("/")

# [FN:root] ----------------------------------------------------
def root():
    return {
        "ok": True,
        "service": "Evolv Training Copilot API",
        "hint": "Try /health, /ask (POST), or /admin/*",
    }
    
@app.post("/debug/retrieve")

# [FN:debug_retrieve] ----------------------------------------------------
def debug_retrieve(payload: AskPayload):
    q = (payload.question or "").strip()
    if not q:
        raise HTTPException(status_code=400, detail="Empty question")
    ctx_docs = retrieve(q, payload.top_k)
    # Return raw texts for inspection
    return {
        "items": [
            {
                "kind": d["kind"],
                "id": d["id"],
                "meta": d["meta"],
                "preview": d["text"][:1200]
            } for d in ctx_docs
        ]
    }

import re


# ==============================================================================
# [S13] RULE-BASED ANSWER EXTRACTOR
# ==============================================================================

# [FN:rule_based_extract] ----------------------------------------------------
def rule_based_extract(q: str, ctx_text: str) -> str | None:
    """If the answer is plainly present in context, return it directly."""
    ql = q.lower()

    # 1) Director of Nursing
    if "director of nursing" in ql or "don" in ql:
        m = re.search(r"(?im)^Director of Nursing:\s*(.+)$", ctx_text)
        if m:
            return f"Director of Nursing: {m.group(1).strip()}"

    # 2) Administrator
    if "administrator" in ql or "admin" in ql:
        m = re.search(r"(?im)^Administrator:\s*(.+)$", ctx_text)
        if m:
            return f"Administrator: {m.group(1).strip()}"

    # 3) Preferred Home Health (return one agency)
    if "preferred home health" in ql or "preferred home health agency" in ql:
        # Try a labeled block first
        block = re.search(
            r"(?is)###\s*Preferred Home Health Agencies\s*(.+?)(?:\n###|\Z)",
            ctx_text
        )
        block_text = block.group(1) if block else ctx_text
        # Look for bullet or plain list lines (stop at blank line)
        lines = [ln.strip(" -•\t") for ln in block_text.splitlines()]
        for ln in lines:
            if not ln.strip():
                break
            # Skip headings or notes; pick the first plausible agency name
            if len(ln) >= 3 and not ln.lower().startswith(("phone:", "address:", "aliases:", "notes:")):
                return ln.strip()

    return None



# ==============================================================================
# [S5.3] HELPERS — Facility detection utilities
# ==============================================================================
# ---------- Facility helpers ----------
import time
import unicodedata


# [FN:_norm] ----------------------------------------------------
def _norm(s: str) -> str:
    """Lowercase, strip, collapse spaces, remove common punctuation for fuzzy contains."""
    s = (s or "")
    s = unicodedata.normalize("NFKC", s).lower()
    # keep letters/numbers/spaces only
    out = []
    for ch in s:
        if ch.isalnum() or ch.isspace():
            out.append(ch)
        elif ch in ("-", "&", "/", "@", ".", "+", "_", ":"):
            out.append(" ")
    s = "".join(out)
    s = " ".join(s.split())
    return s


# ---- NEW: tag-aware token helpers --------------------------------------------
import re
TAG_SPLIT = re.compile(r"[,\|/;]+")
WORD_SPLIT = re.compile(r"[^\w@.+-]+")

def _tokenize(text: str) -> set[str]:
    # expand dictionary *before* normalizing/splitting
    s = normalize_dictionary_query(str(text or ""))
    s = _norm(s)
    return {t for t in WORD_SPLIT.split(s) if t}

def _tag_tokens(tag_str: str) -> set[str]:
    # split tags on commas/semicolons/pipes, then tokenize
    parts = [p.strip() for p in TAG_SPLIT.split(str(tag_str or "")) if p.strip()]
    out = set()
    for p in parts:
        out |= _tokenize(p)
    return out

def _tag_overlap_score(q_tokens: set[str], row_tags: str) -> float:
    """Return 0..1 boost based on how many query tokens appear in Tags."""
    t = _tag_tokens(row_tags)
    if not t:
        return 0.0
    hits = len(q_tokens & t)
    # gentle, non-linear boost: 1 hit helps; 2–3 hits help more; capped at 1.0
    return min(1.0, 0.25 * hits)


# ---- NEW: generic "list intent" + contact-ish row shaping --------------------
LIST_LIKE_RE = re.compile(r"\b(list|show|directory|contacts?|staff|team|employees?)\b", re.I)
ROLE_HINT_RE = re.compile(
    r"\b(admin(?:istrator)?|don|director|manager|supervisor|admissions|marketing|social\s*work(?:er)?|therapy|pt|ot|st|md|medical\s*director|ap|np|rn|lvn|lpn)\b",
    re.I,
)

def _list_intent(q: str) -> bool:
    return LIST_LIKE_RE.search(q or "") is not None

def _row_role_hint(meta: dict) -> str:
    """Make a clean role label (Administrator, Director of Nursing, Medical Director, Admissions, etc.)."""
    rid = str(meta.get("id","") or "")
    qx  = str(meta.get("question","") or "")
    tg  = str(meta.get("tags","") or "")

    # 1) Prefer role tokens from question/tags
    m = ROLE_HINT_RE.search(qx) or ROLE_HINT_RE.search(tg)
    if m:
        raw = m.group(0).lower()
        # normalize common roles
        norm_map = {
            "admin": "Administrator",
            "administrator": "Administrator",
            "don": "Director of Nursing",
            "director of nursing": "Director of Nursing",
            "md": "Medical Director",
            "medical director": "Medical Director",
            "pt": "Physical Therapy",
            "ot": "Occupational Therapy",
            "st": "Speech Therapy",
            "np": "Nurse Practitioner",
            "pa": "Physician Assistant",
            "aprn": "APRN",
            "rn": "RN",
            "lpn": "LPN",
            "lvn": "LVN",
        }
        return norm_map.get(raw, raw.title())

    # 2) id after ":" is often a decent label (e.g., client:Administrator)
    if ":" in rid:
        label = rid.split(":",1)[1].strip()
        if label:
            return label

    # 3) fallback: trimmed question without '?'
    return re.sub(r"\?$","",qx).strip() or rid


# Contact/Directory intent must be explicit (no generic 'name' anymore)
STAFF_ROLE_RE = re.compile(
    r"\b(admin(?:istrator)?|don|director of nursing|director|admissions|marketing|social\s*work(?:er)?|"
    r"therapy|pt|ot|st|md|medical\s*director|np|pa|aprn|rn|lpn|lvn)\b",
    re.I,
)
CONTACT_FIELD_RE = re.compile(r"\b(email|phone|cell|mobile|fax|contact)\b", re.I)
CONTACT_INTENT_RE = re.compile(r"\b(staff|team|directory|contacts?|lead|leadership|department|roles?)\b", re.I)

def _looks_contactish(meta: dict) -> bool:
    """A row counts as directory/contact only if:
       - it clearly names a staff role (Administrator, DON, MD, etc), OR
       - it mentions a contact field (email/phone) AND there is directory/staff intent.
       This prevents 'legal name' rows or generic detail rows from leaking in.
    """
    qx = str(meta.get("question","") or "")
    tg = str(meta.get("tags","") or "")

    # 1) Obvious role mention in question or tags
    if STAFF_ROLE_RE.search(qx) or STAFF_ROLE_RE.search(tg):
        return True

    # 2) Has a contact field AND the row is in a directory/staff context
    if (CONTACT_FIELD_RE.search(qx) or CONTACT_FIELD_RE.search(tg)) and (CONTACT_INTENT_RE.search(qx) or CONTACT_INTENT_RE.search(tg)):
        return True

    return False


def _render_list_from_candidates(fac_name: str, candidates: list[dict]) -> tuple[str, list[str]]:
    """
    Deterministic list render from candidate QA rows.
    Returns (markdown, ids_used). Only uses provided rows; no model paraphrase.
    """
    items, ids = [], []
    for r in candidates:
        if r.get("kind") != "qa":
            continue
        meta = r.get("meta", {}) or {}

        # Skip anything that's not clearly staff/contact-ish
        if not _looks_contactish(meta):
            continue

        q_lower = str(meta.get("question","") or "").strip().lower()
        # Exclude generic facility detail questions from lists
        if q_lower.startswith("what is the legal name") or q_lower.startswith("are details for"):
            continue

        label = _row_role_hint(meta)
        val   = str(meta.get("answer","") or "").strip()
        if not val:
            continue

        items.append((label, val, str(r.get("id",""))))
        ids.append(str(r.get("id","")))

    # de-dup by (label, value)
    seen = set()
    uniq = []
    for label, val, rid in items:
        key = (label.lower(), val.lower())
        if key not in seen:
            seen.add(key)
            uniq.append((label, val, rid))
    if not uniq:
        return "", []
    uniq.sort(key=lambda x: x[0].lower())
    lines = []
    if fac_name:
        lines.append(f"**Directory for {fac_name}**")
    for label, val, _ in uniq:
        lines.append(f"- {label}: {val}")
    return "\n".join(lines), ids
# ------------------------------------------------------------------------------



# [FN:_parse_facility_tags] ----------------------------------------------------
def _parse_facility_tags(tag_str: str) -> list[str]:
    """Return facility names from a tag string using f- prefix."""
    out = []
    for raw in (tag_str or "").split(","):
        t = raw.strip()
        if t.lower().startswith("f-"):
            # keep original after f- as canonical label; also return a normalized key
            name = t[2:].strip()
            if name:
                out.append(name)
    return out

# cache of known facilities mined from CSV
_facility_catalog = {"_mtime": 0.0, "names": [], "norms": []}


# [FN:_refresh_facility_catalog] ----------------------------------------------------
def _refresh_facility_catalog():
    """Load all f- facilities from qa.csv once (or when the CSV changes)."""
    try:
        mtime = os.path.getmtime(CSV_PATH)
    except Exception:
        mtime = 0.0
    if mtime <= _facility_catalog.get("_mtime", 0.0):
        return  # up to date

    df = load_csv_df()
    names = []
    for tags in df.get("tags", []):
        names.extend(_parse_facility_tags(str(tags)))
    # de-dupe while preserving order
    seen = set()
    uniq = []
    for n in names:
        if n not in seen:
            seen.add(n)
            uniq.append(n)

    _facility_catalog["_mtime"] = mtime
    _facility_catalog["names"] = uniq
    _facility_catalog["norms"] = [_norm(n) for n in uniq]


# [FN:_detect_facility_from_question] ----------------------------------------------------
def _detect_facility_from_question(q: str) -> str | None:
    # first pass: alias-driven detect
    maps = _dictionary_maps()
    alias_map = maps["facility_aliases"]  # key(lower) -> canonical name
    qn = _norm(normalize_dictionary_query(q))

    for k, canon in alias_map.items():
        if _norm(k) in qn:
            return canon  # return canonical facility name

    # fallback: the existing catalog scan
    _refresh_facility_catalog()
    best = None
    for name, nn in zip(_facility_catalog["names"], _facility_catalog["norms"]):
        if nn and nn in qn:
            if best is None or len(nn) > len(_norm(best)):
                best = name
    return best


# [FN:_match_facility_in_meta] ----------------------------------------------------
def _match_facility_in_meta(meta: dict, target_name: str | None) -> bool:
    """Check if a QA row's tags include f-<target_name> (lenient match)."""
    if not target_name:
        return False
    tags = str(meta.get("tags", ""))
    facs = _parse_facility_tags(tags)
    if not facs:
        return False
    tn = _norm(target_name)
    for f in facs:
        if _norm(f) == tn:
            return True
    return False


# [FN:_filter_candidates_by_facility] ------------------------------------------
def _filter_candidates_by_facility(candidates: list[dict], fac_name: str | None) -> list[dict]:
    """
    Prefer strict tag-based match (f-<facility>) and fall back to a light
    textual match when tags are missing. Only keeps QA rows.
    """
    if not fac_name:
        return [c for c in candidates if c.get("kind") == "qa"]

    # 1) STRICT: tag match (best case)
    strict = [
        c for c in candidates
        if c.get("kind") == "qa" and _match_facility_in_meta(c.get("meta", {}), fac_name)
    ]
    if strict:
        return strict

    # 2) FALLBACK: fuzzy contains in question/answer/tags if tags are missing
    qn = _norm(fac_name)
    def _has_text_match(meta):
        q = (meta.get("question","") or "") + " " + (meta.get("answer","") or "") + " " + (meta.get("tags","") or "")
        return _norm(q).find(qn) != -1

    return [
        c for c in candidates
        if c.get("kind") == "qa" and _has_text_match(c.get("meta", {}))
    ]

    
# [FN:_facility_details_answer] ----------------------------------------------------
def _facility_details_answer(fac_name: str) -> tuple[str, list[str]]:
    """
    Build a single answer that lists ALL fields for a facility from qa.csv where:
      - section is either "Facility Details" OR "Client Details" (case-insensitive)
      - tags contain an f-<facility_name> (lenient match)
    Returns: (answer_markdown, list_of_ids_used)
    """
    df = load_csv_df()
    if df.empty:
        return "", []

    # Accept both common sections
    sec = df["section"].fillna("").str.strip().str.lower()
    sec_mask = sec.isin({"facility details".lower(), "client details".lower()})
    df = df[sec_mask] if sec_mask.any() else df  # if nothing matches, don't throw everything away

    # --- robust tag matching with abv-aware token subset ---
    import re
    STOPWORDS = {"the","at","of","and","&","for","in","on","to","from"}

    def _tokens(s: str) -> set[str]:
        # expand abbreviations BEFORE normalizing/splitting (so Abv List helps here too)
        s = normalize_abbrev_query(s)
        s = _norm(s)
        toks = [t for t in re.split(r"\s+", s) if t]
        return {t for t in toks if t not in STOPWORDS}

    target_norm   = _norm(fac_name)
    target_tokens = _tokens(fac_name)

    def _match_tags(tag_str: str) -> bool:
        for tag_name in _parse_facility_tags(str(tag_str or "")):
            # 1) exact normalized equality
            if _norm(tag_name) == target_norm:
                return True
            # 2) token-subset: all query tokens appear in the tag tokens
            if target_tokens and target_tokens.issubset(_tokens(tag_name)):
                return True
        return False

    tag_mask = df["tags"].apply(_match_tags)
    sub = df[tag_mask]
    if sub.empty:
        return "", []

    rows, ids = [], []
    for _, r in sub.iterrows():
        rid = str(r.get("id", "")).strip()
        q   = str(r.get("question", "")).strip()
        a   = str(r.get("answer", "")).strip()

        # Label: prefer text after "<slug>:" in id; else fall back to the question
        label = ""
        if ":" in rid:
            label = rid.split(":", 1)[1].strip()
        if not label:
            label = q

        rows.append((label, a))
        ids.append(rid)

    rows.sort(key=lambda x: x[0].lower())

    lines = [f"**Facility:** {fac_name}"]
    for label, val in rows:
        lines.append(f"- {label}: {val}")
    return "\n".join(lines), ids



# [FN:_extract_facility_name_from_query] -----------------------------------------
_FD_PATTERNS = [
    r"(?i)\b(?:provide|show|list|give|send|return)\s+all\s+facility\s+details\s+for\s+(.+)$",
    r"(?i)\bfacility\s+details\s+for\s+(.+)$",
    r"(?i)\ball\s+facility\s+details\b.*\bfor\s+(.+)$",
]

def _extract_facility_name_from_query(q: str) -> str | None:
    """
    Try a few forgiving patterns. If they fail, fall back to the catalog detector.
    Accepts bracketed [Name], quoted "Name", or plain text.
    """
    q = (q or "").strip()
    # Prefer explicit bracketed form [Facility Name]
    m = re.search(r"\[\s*(.+?)\s*\]", q)
    if m and m.group(1).strip():
        return m.group(1).strip()

    # Quoted "Facility Name"
    m = re.search(r'"([^"]+)"', q)
    if m and m.group(1).strip():
        return m.group(1).strip()

    # Pattern-based “... for NAME”
    for pat in _FD_PATTERNS:
        m = re.search(pat, q)
        if m and m.group(1).strip():
            return m.group(1).strip()

    # Fallback to catalog-driven fuzzy detector
    return _detect_facility_from_question(q)



# [FN:upsert_facility_and_children] --------------------------------------------
def upsert_facility_and_children(payload: dict):
    """
    payload keys (from TEST Facility_Details.html -> collectValues()):
      - facility_name (from header or URL f=...)
      - legal_name, corporate_group, address_line1, address_line2, city, state, zip, county
      - avg_dcs, short_beds, ltc_beds, outpatient_pt
      - emr, emr_other, pt_emr
      - orders, orders_other
      - additional_services: [str, ...]
      - insurance_plans:    [str, ...]
      - contacts:           [{type,name,email,phone,pref}, ...]
      - community_partners: [{type,name,ins_only,insurance}, ...]

    Strategy:
      1) derive facility_id (prefer payload.facility_id or slug of facility_name)
      2) upsert into facilities
      3) replace children (delete+bulk insert)
    """
    import re, datetime, json
    conn = _db()
    cur  = conn.cursor()

    # --- 1) facility_id ---
    facility_name = (payload.get("facility_name") or payload.get("legal_name") or "").strip()
    if not facility_name:
        raise HTTPException(status_code=400, detail="facility_name (or legal_name) is required")

    # allow caller to pass a stable facility_id; else slugify from name
    facility_id = (payload.get("facility_id") or "").strip()
    if not facility_id:
        s = re.sub(r'[^a-z0-9-]+', '-', facility_name.lower().strip())
        s = re.sub(r'-+', '-', s).strip('-') or "facility"
        facility_id = s

    # --- 2) upsert single-value columns into facilities ---
    cols = [
        "facility_id","facility_name","legal_name","corporate_group",
        "address_line1","address_line2","city","state","zip","county",
        "avg_dcs","short_beds","ltc_beds","outpatient_pt",
        "emr","emr_other","pt_emr","orders","orders_other",
        "raw_json"
    ]
    row = {
        "facility_id": facility_id,
        "facility_name": facility_name,
        "legal_name": payload.get("legal_name",""),
        "corporate_group": payload.get("corporate_group",""),
        "address_line1": payload.get("address_line1",""),
        "address_line2": payload.get("address_line2",""),
        "city": payload.get("city",""),
        "state": payload.get("state",""),
        "zip": payload.get("zip",""),
        "county": payload.get("county",""),
        "avg_dcs": payload.get("avg_dcs",""),
        "short_beds": payload.get("short_beds",""),
        "ltc_beds": payload.get("ltc_beds",""),
        "outpatient_pt": payload.get("outpatient_pt",""),
        "emr": payload.get("emr",""),
        "emr_other": payload.get("emr_other",""),
        "pt_emr": payload.get("pt_emr",""),
        "orders": payload.get("orders",""),
        "orders_other": payload.get("orders_other",""),
        "raw_json": json.dumps(payload, ensure_ascii=False)
    }

    cur.execute("""
    INSERT INTO facilities(
      facility_id, facility_name, legal_name, corporate_group,
      address_line1, address_line2, city, state, zip, county,
      avg_dcs, short_beds, ltc_beds, outpatient_pt,
      emr, emr_other, pt_emr, orders, orders_other,
      raw_json, created_at, updated_at
    )
    VALUES (:facility_id, :facility_name, :legal_name, :corporate_group,
      :address_line1, :address_line2, :city, :state, :zip, :county,
      :avg_dcs, :short_beds, :ltc_beds, :outpatient_pt,
      :emr, :emr_other, :pt_emr, :orders, :orders_other,
      :raw_json, datetime('now'), datetime('now')
    )
    ON CONFLICT(facility_id) DO UPDATE SET
      facility_name=excluded.facility_name,
      legal_name=excluded.legal_name,
      corporate_group=excluded.corporate_group,
      address_line1=excluded.address_line1,
      address_line2=excluded.address_line2,
      city=excluded.city,
      state=excluded.state,
      zip=excluded.zip,
      county=excluded.county,
      avg_dcs=excluded.avg_dcs,
      short_beds=excluded.short_beds,
      ltc_beds=excluded.ltc_beds,
      outpatient_pt=excluded.outpatient_pt,
      emr=excluded.emr,
      emr_other=excluded.emr_other,
      pt_emr=excluded.pt_emr,
      orders=excluded.orders,
      orders_other=excluded.orders_other,
      raw_json=excluded.raw_json,
      updated_at=datetime('now');
    """, row)

    # --- 3) replace children ---
    addl = payload.get("additional_services") or []
    plans = payload.get("insurance_plans") or []
    contacts = payload.get("contacts") or []
    partners = payload.get("community_partners") or []

    cur.execute("DELETE FROM facility_additional_services WHERE facility_id = ?", (facility_id,))
    cur.execute("DELETE FROM facility_insurance_plans   WHERE facility_id = ?", (facility_id,))
    cur.execute("DELETE FROM facility_contacts          WHERE facility_id = ?", (facility_id,))
    cur.execute("DELETE FROM facility_partners          WHERE facility_id = ?", (facility_id,))

    if addl:
        cur.executemany(
            "INSERT INTO facility_additional_services(facility_id, service) VALUES (?, ?)",
            [(facility_id, str(s).strip()) for s in addl if str(s).strip()]
        )
    if plans:
        cur.executemany(
            "INSERT INTO facility_insurance_plans(facility_id, plan) VALUES (?, ?)",
            [(facility_id, str(p).strip()) for p in plans if str(p).strip()]
        )
    if contacts:
        cur.executemany(
            "INSERT INTO facility_contacts(facility_id, type, name, email, phone, pref) VALUES (?, ?, ?, ?, ?, ?)",
            [
              (facility_id,
               str(c.get("type","")).strip(),
               str(c.get("name","")).strip(),
               str(c.get("email","")).strip(),
               str(c.get("phone","")).strip(),
               str(c.get("pref","")).strip())
              for c in contacts
              if (c and str(c.get("type","")).strip() and str(c.get("name","")).strip())
            ]
        )
    if partners:
        cur.executemany(
            "INSERT INTO facility_partners(facility_id, type, name, ins_only, insurance) VALUES (?, ?, ?, ?, ?)",
            [
              (facility_id,
               str(p.get("type","")).strip(),
               str(p.get("name","")).strip(),
               str(p.get("ins_only","")).strip(),
               str(p.get("insurance","")).strip())
              for p in partners
              if (p and str(p.get("type","")).strip() and str(p.get("name","")).strip())
            ]
        )

    conn.commit()
    conn.close()
    return facility_id




# --------------------------------------





# ==============================================================================
# [S14] ROUTE — /ask (Main answer endpoint)
# ==============================================================================
@app.post("/ask")
# [FN:ask] ----------------------------------------------------
def ask(payload: AskPayload):
    """
    Main Q&A endpoint used by the HTML Ask page.

    It:
      - Normalizes the question
      - Applies optional topic hints (Sensys / List/Group / etc.)
      - Tries a few fast-paths (contact lookup, facility details, list/group)
      - Falls back to retrieval + model answer
    """
    q = normalize_dictionary_query((payload.question or "").strip())
    q_norm = normalize_abbrev_query(q)
    if not q:
        raise HTTPException(status_code=400, detail="Empty question")

    # Accept either 'section' or 'section_hint' from the client
    section_pref = (payload.section or payload.section_hint or "").strip() or None

    # Topic flavor hints (multi-select via comma list)
    topics_lower = [t.strip().lower() for t in (section_pref or "").split(",") if t.strip()]
    HOWTO_MODE = any(t in ("sensys", "how-to", "how to") for t in topics_lower)
    LIST_MODE  = any(t in ("list/group", "list", "group") for t in topics_lower)

    # ---- NEW: direct contact lookup (SQLite) ---------------------------------
    # If the user asks for staff by role for a given facility, answer from DB first.
    ql = q.lower()
    fac_guess = _infer_facility_from_question(q) or _detect_facility_from_question(q) or ""
    role_tokens = [
        ("administrator", ["administrator", "admin"]),
        ("director of nursing", ["director of nursing", "don"]),
        ("admissions", ["admissions"]),
        ("marketing", ["marketing"]),
        ("social services", ["social services", "social worker"]),
        ("therapy", ["therapy"]),
        ("medical director", ["medical director"]),
        ("physician/np", ["physician", "doctor", "md", "np"]),
    ]
    role_hit = None
    for label, keys in role_tokens:
        if any(k in ql for k in keys):
            role_hit = label
            break

    if fac_guess and role_hit:
        got = _fetch_contact_from_sqlite(fac_guess, role_hit)
        if got:
            role_label, value_text = got
            log_id = None
            try:
                log_id = _append_user_qa_log("Ask", q, f"{role_label}: {value_text}")
            except Exception as e:
                print("log append error (contact fast-path):", e)

            payload_out = {
                "answer": f"{role_label}: {value_text}",
                "sources": [f"DB:facility_contacts:{fac_guess}:{role_label}"],
                "section_used": section_pref or "",
                "log_id": log_id or None,
            }
            try:
                cache_key = (normalize_abbrev_query(q), section_pref or "")
                _ask_cache_put(cache_key, payload_out)
            except Exception:
                pass
            return payload_out
    # --------------------------------------------------------------------------

    # === Contact lookup (DB-first, facility + role) ===========================
    fac_name = _extract_facility_name_from_query(q) or _detect_facility_from_question(q)
    role = _extract_contact_role(q)
    if fac_name and role:
        md, meta_used = _contact_answer_from_db(fac_name, role)
        if md:
            log_id = None
            try:
                log_id = _append_user_qa_log("Contacts", q, md)
            except Exception as e:
                print("log append error (contacts DB):", e)

            payload_out = {
                "answer": md,
                "sources": [f"DB:facility_contacts:{meta_used.get('facility_id','')}"],
                "section_used": "Contacts",
                "log_id": log_id or None,
            }
            cache_key = (q_norm, section_pref or "")
            _ask_cache_put(cache_key, payload_out)
            return payload_out
    # --------------------------------------------------------------------------

    # ---- List/Group early-route (topic selected) -----------------------------
    if LIST_MODE:
        lg = build_list_group_answer(q)
        if lg and lg[0]:
            log_id = None
            try:
                log_id = _append_user_qa_log("List/Group", q, lg[0])
            except Exception as e:
                print("log append error (List/Group):", e)

            payload_out = {
                "answer": lg[0],
                "sources": [],
                "section_used": "List/Group",
                "log_id": log_id or None,
            }
            try:
                cache_key = (normalize_abbrev_query(q), "List/Group")
                _ask_cache_put(cache_key, payload_out)
            except Exception as e:
                print("ask cache skip:", e)
            return payload_out
    # --------------------------------------------------------------------------

    # Optional: adjust system prompt for Sensys how-to style
    system_prompt = SYSTEM_PROMPT  # start from the global base for this request only
    if HOWTO_MODE:
        system_prompt = (
            SYSTEM_PROMPT
            + "\n- For Sensys topics, answer in numbered, concise step-by-step instructions."
        )

    # === Special case: Facility Details → go straight to SQLite ===============
    fd_name = _extract_facility_name_from_query(q) or _detect_facility_from_question(q)
    fd_topic_selected = bool(section_pref and "facility details" in (section_pref or "").lower())
    if fd_name and (fd_topic_selected or "facility details" in _norm(q)):
        md, meta_used = _facility_details_answer_from_db(fd_name)
        if md:
            log_id = None
            try:
                log_id = _append_user_qa_log("Facility Details", q, md)
            except Exception as e:
                print("log append error (facility details DB):", e)

            payload_out = {
                "answer": md,
                "sources": [f"DB:facilities:{meta_used.get('facility_id','')}"],
                "section_used": "Facility Details",
                "log_id": log_id or None,
            }
            cache_key = (normalize_abbrev_query(q), "Facility Details")
            _ask_cache_put(cache_key, payload_out)
            return payload_out
    # ==========================================================================

    # --- TTL cache: same question + section -----------------------------------
    cache_key = (q_norm, section_pref or "")
    cached = _ask_cache_get(cache_key)
    if cached:
        return cached

    # Use section-aware retrieval
    ctx_docs = retrieve(q_norm, payload.top_k, section=section_pref)

    # If it's a list/directory/staff/team intent, pull a larger pool once
    try:
        if "_list_intent" in globals() and _list_intent(q):
            ctx_docs = retrieve(q_norm, max(payload.top_k, 60), section=section_pref)
    except Exception as e:
        print("list intent pool bump skipped:", e)

    # Debug: return top candidates with component scores
    if payload.debug:
        preview = []
        for d in ctx_docs[:30]:
            if d["kind"] != "qa":
                continue
            meta = d.get("meta", {}) or {}
            preview.append({
                "id": d.get("id", ""),
                "kw": round(float(d.get("_kw", 0.0)), 3),
                "tagb": round(float(d.get("_tagb", 0.0)), 3),
                "pre": round(float(d.get("_pre", 0.0)), 3),
                "score": round(float(d.get("score", 0.0)), 3),
                "section": meta.get("section", ""),
                "tags": (meta.get("tags", "") or "")[:160],
                "question": (meta.get("question", "") or "")[:160],
            })
        return {
            "answer": "Debug only — top candidates with scores",
            "debug_candidates": preview,
            "sources": [f"Q&A:{d.get('id', '')}" for d in ctx_docs[:30] if d["kind"] == "qa"],
            "section_used": section_pref or "",
        }

    # ---- generic "list mode" (directory/contacts/staff/team) -----------------
    fac_name = _extract_facility_name_from_query(q) or _detect_facility_from_question(q)
    if "_list_intent" in globals() and _list_intent(q):
        scoped = _filter_candidates_by_facility(ctx_docs, fac_name)
        md, used_ids = _render_list_from_candidates(fac_name, scoped)
        if md:
            log_id = None
            try:
                log_id = _append_user_qa_log(section_pref or "", q, md)
            except Exception as e:
                print("log append error (list mode):", e)
            return {
                "answer": md,
                "sources": [f"Q&A:{rid}" for rid in used_ids],
                "section_used": section_pref or "",
                "log_id": log_id or None,
            }
    # --------------------------------------------------------------------------

    # Build context & readable sources
    ctx_block_parts = []
    readable_sources = []
    for d in ctx_docs:
        if d["kind"] == "qa":
            q_text = d["meta"].get("question", "")
            a_text = d["meta"].get("answer", "")
            t_text = d["meta"].get("tags", "")
            ctx_block_parts.append(f"[Q&A {d['id']}]\nQ: {q_text}\nA: {a_text}\nTags: {t_text}")
            readable_sources.append(f"Q&A:{d['id']}")
        else:
            src = d["meta"].get("source", "doc")
            ch = d["meta"].get("chunk", "?")
            ctx_block_parts.append(f"[DOC {src} #{ch}]\n{d['text'][:2000]}")
            readable_sources.append(f"DOC:{src}#{ch}")

    ctx_block = "\n\n".join(ctx_block_parts) if ctx_block_parts else ""

    # Try a quick deterministic extraction before calling the model
    rb_answer = rule_based_extract(q, ctx_block)
    if rb_answer:
        return {"answer": rb_answer, "sources": readable_sources}

    used_ids = ", ".join(readable_sources) if readable_sources else "none"
    user_msg = (
        f"User question: {q}\n\n"
        f"Context Q&A/Docs:\n{ctx_block if ctx_block else '(no relevant context found)'}\n\n"
        "Return:\n"
        "- Direct answer ONLY.\n"
        "- DO NOT include any citations, source IDs, or brackets like [Sources: ...] in the answer.\n"
        f"- (For reference only, sources are: {used_ids})"
    )

    try:
        resp = client.responses.create(
            model="gpt-4o-mini",
            input=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_msg},
            ],
            max_output_tokens=350,
        )
        answer = (resp.output_text or "").strip()
        print("DEBUG raw output_text repr:", repr(resp.output_text))

        if not answer:
            if ctx_docs and ctx_docs[0]["kind"] == "qa":
                top = ctx_docs[0]
                answer = f"{top['meta'].get('answer', '')}"
            else:
                answer = "I couldn't find anything relevant. Try rephrasing or add more entries."

        log_id = None
        try:
            log_id = _append_user_qa_log(section_pref or "", q, answer)
        except Exception as e:
            print("log append error:", e)

        payload_out = {
            "answer": answer,
            "sources": readable_sources,
            "section_used": section_pref or "",
            "log_id": log_id or None,
        }
        _ask_cache_put(cache_key, payload_out)
        return payload_out
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"OpenAI error: {e}")


def _extract_contact_role(q: str) -> str | None:
    """
    Pull a contact role from a free-text question.
    Returns the normalized label (e.g., 'Administrator', 'Director of Nursing', etc.)
    or None if no role is detected.
    """
    ql = (q or "").lower()
    role_tokens = [
        ("Administrator", ["administrator", "admin"]),
        ("Director of Nursing", ["director of nursing", "don"]),
        ("Admissions", ["admissions"]),
        ("Marketing", ["marketing"]),
        ("Social Services", ["social services", "ss"]),
        ("Therapy", ["therapy", "pt", "ot", "st"]),
        ("Nursing", ["nursing"]),
        ("Physician", ["physician", "doctor", "md", "do"]),
        ("Medical Director", ["medical director"]),
    ]
    for label, keys in role_tokens:
        if any(k in ql for k in keys):
            return label
    return None



# ==============================================================================
# [S15] ROUTES: Q&A Admin — List/Add/Update/Delete/Upload/Reindex
# ==============================================================================
# -------- Admin: list/add/update/delete/upload/reindex --------
@app.get("/admin/list")

# [FN:admin_list] ----------------------------------------------------
def admin_list(request: Request):
    require_admin(request)
    df = load_qa_df_unified()
    return {"rows": df.to_dict(orient="records"), "count": len(df)}


# [FN:admin_top_tags] ----------------------------------------------------
@app.get("/admin/top_tags")
def admin_top_tags(request: Request, limit: int = 5):
    require_admin(request)
    df = load_qa_df_unified()
    tags_series = df["tags"].dropna().astype(str)

    import re
    freq = {}
    for s in tags_series:
        # split on commas/pipes/semicolons
        parts = re.split(r"[,\|;]", s)
        for t in parts:
            tag = (t or "").strip()
            if not tag:
                continue
            # ignore facility tags in popularity counts
            if tag.lower().startswith("f-"):
                continue
            key = tag.lower()
            freq[key] = freq.get(key, 0) + 1

    top = sorted(freq.items(), key=lambda x: x[1], reverse=True)[:limit]
    return {"tags": [t for t, _ in top]}



@app.post("/admin/add")
# [FN:admin_add] ----------------------------------------------------
def admin_add(
    request: Request,
    # NOTE: ID is no longer accepted from the client — server generates it
    question: str = Form(...),
    answer: str = Form(...),
    tags: str = Form(""),
    section: str = Form(""),
    topics: str = Form(""),  # keep for UI; falls back to section if blank
):
    # Auth as you already do
    require_admin(request)

    # Basic validation
    if not (question or "").strip():
        raise HTTPException(status_code=400, detail="Question is required")
    if not (answer or "").strip():
        raise HTTPException(status_code=400, detail="Answer is required")

    # Generate unique ID and write atomically within a process lock
    with QA_LOCK:
        df = load_qa_df_unified()
        rows = df.to_dict(orient="records")
        new_id = _next_id_from_rows(rows)  # e.g., Q001, Q002, ...

        new_row = {
            "id":       new_id,
            "section":  (section or "").strip(),
            "topics":   (topics or section).strip(),  # back-compat default
            "question": question.strip(),
            "answer":   answer.strip(),
            "tags":     tags.strip(),
        }

        # append and persist
        df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
        if USE_SQLITE_WRITE:
            # primary write → SQLite
            upsert_qa_sqlite(new_row)
            # optional CSV mirror (turn off later)
            if MIRROR_SQLITE_WRITES:
                save_csv_df(df)
        else:
            # primary write → CSV (old behavior)
            save_csv_df(df)
            # mirror to SQLite (as in Step 3.3)
            if MIRROR_SQLITE_WRITES:
                try:
                    upsert_qa_sqlite(new_row)
                except Exception as e:
                    print("mirror add->sqlite failed:", e)

    # Try to refresh the index (non-fatal on failure)
    try:
        rebuild_index_from_csv()
    except Exception as e:
        print("Reindex after add failed:", e)

    # Return the created row so the UI can show the server-issued ID immediately
    return {"ok": True, "row": new_row, "count": int(len(df))}


@app.post("/admin/update")

# [FN:admin_update] ----------------------------------------------
def admin_update(
    request: Request,
    id: str = Form(...),
    question: str = Form(None),
    answer: str = Form(None),
    tags: str = Form(None),
    section: str = Form(None),
    topics: str = Form(None),  # NEW
):
    require_admin(request)
    df = load_qa_df_unified()
    mask = (df["id"] == str(id))
    if not mask.any():
        raise HTTPException(status_code=404, detail=f"id '{id}' not found")
    if question is not None:
        df.loc[mask, "question"] = question.strip()
    if answer is not None:
        df.loc[mask, "answer"] = answer.strip()
    if tags is not None:
        df.loc[mask, "tags"] = (tags or "").strip()
    if section is not None:
        df.loc[mask, "section"] = (section or "").strip()
    if topics is not None:
        df.loc[mask, "topics"] = (topics or "").strip()
    save_csv_df(df)
    # NEW (Step 3.3): mirror update to SQLite
    row_now = df.loc[mask].iloc[0].to_dict()
    row_now.setdefault("topics", row_now.get("section",""))

    if USE_SQLITE_WRITE:
        # primary write → SQLite
        upsert_qa_sqlite(row_now)
        # optional CSV mirror
        if MIRROR_SQLITE_WRITES:
            save_csv_df(df)
    else:
        # primary write → CSV
        save_csv_df(df)
        # mirror to SQLite
        if MIRROR_SQLITE_WRITES:
            try:
                upsert_qa_sqlite(row_now)
            except Exception as e:
                print("mirror update->sqlite failed:", e)
    return {"ok": True}

@app.post("/admin/delete")

# [FN:admin_delete] ----------------------------------------------------
def admin_delete(request: Request, id: str = Form(...)):
    require_admin(request)
    df = load_qa_df_unified()
    before = len(df)
    df = df[df["id"] != str(id)]
    if len(df) == before:
        raise HTTPException(status_code=404, detail=f"id '{id}' not found")
    save_csv_df(df)
    if USE_SQLITE_WRITE:
        # primary write → SQLite
        delete_qa_sqlite(str(id))
        # optional CSV mirror
        if MIRROR_SQLITE_WRITES:
            save_csv_df(df)
    else:
        # primary write → CSV
        save_csv_df(df)
        # mirror to SQLite
        if MIRROR_SQLITE_WRITES:
            try:
                delete_qa_sqlite(str(id))
            except Exception as e:
                print("mirror delete->sqlite failed:", e)
    return {"ok": True, "count": len(df)}

@app.post("/admin/upload_csv")

# [FN:admin_upload_csv] ----------------------------------------------------
def admin_upload_csv(request: Request, file: UploadFile = File(...)):
    require_admin(request)
    content = file.file.read()
    try:
        df = pd.read_csv(io.BytesIO(content)).fillna("")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid CSV: {e}")

    required = {"id", "question", "answer", "tags"}
    # Accept section if present; create if missing
    if not required.issubset(df.columns):
        raise HTTPException(status_code=400, detail=f"CSV must include columns: {sorted(required)}")
    if "section" not in df.columns:
        df["section"] = ""

    df["id"] = df["id"].astype(str)
    save_csv_df(df)
    try:
        rebuild_index_from_csv()
    except Exception as e:
        print("Reindex after upload failed:", e)
    return {"ok": True, "count": len(df)}


# ==============================================================================
# [S15A] ADMIN — one-click CSV → SQLite migration (idempotent)
# ==============================================================================
@app.post("/admin/migrate_to_sqlite")
def admin_migrate_to_sqlite(request: Request):
    require_admin(request)
    _init_db()
    conn = _db()
    cur  = conn.cursor()

    # --- 1) qa.csv -> qa table ---
    qa_df = load_csv_df()  # already normalizes topics/section
    qa_rows = qa_df.to_dict(orient="records")
    migrated_qa = 0
    for r in qa_rows:
        cur.execute("""
            INSERT INTO qa (id, section, topics, question, answer, tags)
            VALUES (?, ?, ?, ?, ?, ?)
            ON CONFLICT(id) DO UPDATE SET
              section=excluded.section,
              topics =excluded.topics,
              question=excluded.question,
              answer=excluded.answer,
              tags   =excluded.tags
        """, (
            str(r.get("id","")),
            str(r.get("section","")),
            str(r.get("topics","")),
            str(r.get("question","")),
            str(r.get("answer","")),
            str(r.get("tags","")),
        ))
        migrated_qa += 1

    # --- 2) abbreviations.csv -> abbreviations table ---
    try:
        abv_rows = _read_abv_rows()
    except Exception:
        abv_rows = []
    migrated_abv = 0
    for r in abv_rows:
        cur.execute("""
            INSERT INTO abbreviations (abbr, meaning, notes)
            VALUES (?, ?, ?)
            ON CONFLICT(abbr) DO UPDATE SET
              meaning=excluded.meaning,
              notes  =excluded.notes
        """, (
            str(r.get("abbr","")).strip(),
            str(r.get("meaning","")).strip(),
            str(r.get("notes","")).strip(),
        ))
        migrated_abv += 1

    # --- 3) user_qa_log.csv -> user_qa_log table ---
    try:
        log_rows = _read_log_rows()
    except Exception:
        log_rows = []
    migrated_log = 0
    for r in log_rows:
        # Use INSERT OR REPLACE to keep explicit ids; table has INTEGER PRIMARY KEY
        cur.execute("""
            INSERT OR REPLACE INTO user_qa_log
                (id, ts, section, q, a, promoted, a_quality)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (
            int(r.get("id", 0) or 0),
            str(r.get("ts","")),
            str(r.get("section","")),
            str(r.get("q","")),
            str(r.get("a","")),
            1 if str(r.get("promoted","")).strip() in ("1","true","True") else 0,
            str(r.get("a_quality","")),
        ))
        migrated_log += 1

    conn.commit()
    conn.close()
    return {
        "ok": True,
        "qa_rows": migrated_qa,
        "abbrev_rows": migrated_abv,
        "log_rows": migrated_log,
        "db_path": DB_PATH
    }




# ==============================================================================
# [S16] BULK INGEST FROM DOCUMENT (PDF/DOCX/TXT)
# ==============================================================================
# ========= BULK INGEST FROM DOCUMENT =========
from pypdf import PdfReader
from docx import Document as DocxDocument
import json
import re


# [FN:_extract_text_from_bytes] ----------------------------------------------------
def _extract_text_from_bytes(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(data))
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                parts.append("")
        return "\n\n".join(parts).strip()
    elif name.endswith(".docx"):
        doc = DocxDocument(io.BytesIO(data))
        return "\n".join([p.text for p in doc.paragraphs]).strip()
    else:
        # default: treat as text
        try:
            return data.decode("utf-8", errors="ignore")
        except Exception:
            return ""


# [FN:_chunk_text] ----------------------------------------------------
def _chunk_text(text: str, max_chars: int = 9000) -> list[str]:
    text = text.strip()
    if len(text) <= max_chars:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        # try to break on paragraph boundary
        cut = text.rfind("\n", start, end)
        if cut == -1 or cut <= start + 2000:  # avoid tiny slices
            cut = end
        chunks.append(text[start:cut].strip())
        start = cut
    return [c for c in chunks if c]


# [FN:_parse_json_loose] ----------------------------------------------------
def _parse_json_loose(s: str):
    """
    Try very hard to parse JSON:
    - strip ```json fences
    - normalize smart quotes
    - extract the first balanced {...} block
    - remove trailing commas before } or ]
    """
    if not s:
        return None

    t = s.strip()

    # 1) Remove code fences like ```json ... ```
    if t.startswith("```"):
        # remove leading ```json or ```
        t = re.sub(r"^```(?:json)?\s*", "", t, flags=re.IGNORECASE)
        # remove trailing ```
        t = re.sub(r"\s*```$", "", t)

    # 2) Normalize smart quotes (common in some model outputs)
    t = (t
         .replace("\u201c", '"')
         .replace("\u201d", '"')
         .replace("\u2018", "'")
         .replace("\u2019", "'"))

    # 3) First, try direct JSON
    try:
        return json.loads(t)
    except Exception:
        pass

    # 4) Extract the first balanced { ... } block
    start = t.find("{")
    if start != -1:
        depth = 0
        end = None
        for i, ch in enumerate(t[start:], start=start):
            if ch == "{":
                depth += 1
            elif ch == "}":
                depth -= 1
                if depth == 0:
                    end = i + 1
                    break
        if end:
            cand = t[start:end]
            # 5) Remove trailing commas like ",}" or ",]"
            cand = re.sub(r",(\s*[}\]])", r"\1", cand)
            try:
                return json.loads(cand)
            except Exception:
                # One more attempt: collapse repeated commas
                cand2 = re.sub(r",\s*,", ",", cand)
                try:
                    return json.loads(cand2)
                except Exception:
                    return None

    return None


# [FN:_harvest_qas_from_string] ----------------------------------------------------
def _harvest_qas_from_string(raw: str, default_section: str, default_tags: str) -> list[dict]:
    """
    Last-resort parser: scans the raw model output (string) and tries to pull
    objects that look like {"section": "...", "question": "...", "answer": "...", "tags": "..."}.
    Tolerant of minor formatting issues.
    """
    if not raw:
        return []
    txt = raw

    # Normalize common quote types
    txt = (txt
           .replace("\u201c", '"').replace("\u201d", '"')
           .replace("\u2018", "'").replace("\u2019", "'"))

    # If it's wrapped in ```json fences, strip them
    if txt.strip().startswith("```"):
        txt = re.sub(r"^```(?:json)?\s*", "", txt.strip(), flags=re.IGNORECASE)
        txt = re.sub(r"\s*```$", "", txt)

    # Try to locate the qas array region first
    m = re.search(r'"qas"\s*:\s*\[', txt)
    start = m.start() if m else 0
    # From start, walk to find the matching closing ] at depth 1 of [ ]
    depth = 0
    end = len(txt)
    for i, ch in enumerate(txt[start:], start=start):
        if ch == '[':
            depth += 1
        elif ch == ']':
            depth -= 1
            if depth == 0:
                end = i + 1
                break
    region = txt[start:end]

    # Now scan the region for top-level {...} objects and pick fields with regex
    out = []
    i = 0
    while True:
        o = region.find("{", i)
        if o == -1:
            break
        depth = 0
        close = None
        for j, ch in enumerate(region[o:], start=o):
            if ch == "{":
                depth += 1
            elif ch == "}":
                depth -= 1
                if depth == 0:
                    close = j + 1
                    break
        if not close:
            break
        obj_txt = region[o:close]

        # Remove trailing commas inside objects (",}")
        obj_txt = re.sub(r",\s*}", "}", obj_txt)

        # Extract fields (order-agnostic, tolerant of spacing)
        def grab(key):
            m = re.search(rf'"{key}"\s*:\s*"(.*?)"', obj_txt, flags=re.DOTALL)
            return (m.group(1).strip() if m else "")

        section = grab("section") or (default_section or "")
        question = grab("question")
        answer = grab("answer")
        tags = grab("tags") or (default_tags or "")

        if question and answer:
            out.append({
                "section": section,
                "question": question,
                "answer": answer,
                "tags": tags
            })

        i = close

    return out




# [FN:_generate_qas_from_text] ----------------------------------------------------
def _generate_qas_from_text(text: str, default_section: str, default_tags: str, thoroughness: str = "high") -> list[dict]:
    """
    Create thorough Q&A pairs from text using the Responses API.
    We enforce JSON via instructions (since this SDK doesn't support response_format here)
    and then parse defensively.
    Returns: list of {section, question, answer, tags}
    """
    sys_msg = (
        "You are a data extraction assistant. Read the provided text and produce MANY specific, high-quality "
        "training Q&A pairs. Each Q&A should be concise, practical, and unambiguous. Avoid duplicates. "
        "Prefer one fact per Q&A. Do not invent content not supported by the text."
    )

    user_prompt = f"""
Return ONLY valid JSON with this exact shape (no prose, no markdown, no extra keys):

{{
  "qas": [
    {{"section":"string","question":"string","answer":"string","tags":"comma,separated"}}
  ]
}}

Guidelines:
- Be THOROUGH: extract as many distinct, useful Q&As as the text truly supports.
- Use the provided default section "{default_section}" unless the text makes a different section obvious.
- Put short, relevant tags (comma-separated).
- Absolutely do not include any text before or after the JSON.

TEXT (may be chunked):
{text}
""".strip()

    resp = client.responses.create(
        model="gpt-5-mini",
        input=[
            {"role": "system", "content": sys_msg},
            {"role": "user", "content": user_prompt},
        ],
        # keep output long enough for lots of Q&As
        max_output_tokens=2200 if thoroughness == "high" else 1200,
        # NOTE: no temperature (unsupported) and no response_format (unsupported in this SDK)
    )

    raw = (resp.output_text or "").strip()
    # Debugging aid in server logs (optional):
    print(f"[ingest] raw len={len(raw)} preview={(raw[:160] + '...') if len(raw) > 160 else raw}")

    data = _parse_json_loose(raw) or {}
    qas = data.get("qas") or []

    cleaned = []
    for q in qas:
        question = (q.get("question") or "").strip()
        answer   = (q.get("answer") or "").strip()
        section  = (q.get("section") or default_section or "").strip()
        tags     = (q.get("tags") or default_tags or "").strip()
        if question and answer:
            cleaned.append({
                "section": section,
                "question": question,
                "answer": answer,
                "tags": tags,
            })

    # Fallback: harvest directly from raw string if JSON parse produced nothing
    if not cleaned:
        harvested = _harvest_qas_from_string(raw, default_section, default_tags)
        if harvested:
            print(f"[ingest] fallback harvested {len(harvested)} QAs from raw")
            cleaned = harvested

    print(f"[ingest] produced {len(cleaned)} QAs from chunk (len={len(text)})")
    return cleaned




# [FN:_next_id] ----------------------------------------------------
def _next_id(existing_ids: set[str], prefix: str = "Q") -> str:
    """
    Generate the next unique ID like Q001, Q002... using existing IDs in the CSV.
    """
    nums = []
    for _id in existing_ids:
        m = re.match(rf"{re.escape(prefix)}(\d+)$", _id)
        if m:
            nums.append(int(m.group(1)))
    nxt = (max(nums) + 1) if nums else 1
    return f"{prefix}{nxt:03d}"

@app.post("/admin/ingest_doc")

# [FN:admin_ingest_doc] ----------------------------------------------------
def admin_ingest_doc(request: Request,
                     file: UploadFile = File(...),
                     section: str = Form(""),
                     tags: str = Form(""),
                     id_prefix: str = Form("Q"),
                     thoroughness: str = Form("high"),
                     reindex: int = Form(1)):
    """
    Upload a PDF/DOCX/TXT. Extract text, generate thorough Q&A pairs with GPT,
    append them to the CSV, and (optionally) rebuild the QA index.
    """
    require_admin(request)

    # Read and extract text
    content = file.file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Empty file")
    text = _extract_text_from_bytes(file.filename or "", content)
    # If there is virtually no text, it's likely a scanned PDF (image-only).
    if len(text.strip()) < 30:
        return {
            "ok": True,
            "created": 0,
            "ids": [],
            "note": "No extractable text found. If this is a scanned PDF, run OCR first (e.g., export as text/searchable PDF) or upload a DOCX/TXT."
        }
    if not text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from file")

    # Chunk for large documents, generate per chunk, then merge
    chunks = _chunk_text(text, max_chars=9000)
    all_qas = []
    for ch in chunks:
        qas = _generate_qas_from_text(ch, section, tags, thoroughness=thoroughness)
        all_qas.extend(qas)

    if not all_qas:
        return {
            "ok": True,
            "created": 0,
            "ids": [],
            "note": (
                "The model responded, but I couldn’t parse any Q&As. "
                "This is usually formatting (fences/quotes/commas). "
                "I’ve now hardened the parser—please try again."
            ),
        }


    # Load CSV, assign IDs, append, save
    df = load_csv_df()
    existing = set(df["id"].astype(str).tolist())
    new_rows = []
    created_ids = []
    for q in all_qas:
        new_id = _next_id(existing, prefix=(id_prefix or "Q"))
        existing.add(new_id)
        new_rows.append({
            "id": new_id,
            "section": q.get("section", section),
            "question": q["question"],
            "answer": q["answer"],
            "tags": q.get("tags", tags),
        })
        created_ids.append(new_id)

    if new_rows:
        df2 = pd.concat([df, pd.DataFrame(new_rows)], ignore_index=True)
        save_csv_df(df2)

    # Optionally reindex the QA collection so results are immediately searchable
    if reindex:
        try:
            rebuild_index_from_csv()
        except Exception as e:
            # Don't fail the whole op if reindex hiccups
            pass

    return {"ok": True, "created": len(new_rows), "ids": created_ids[:100], "truncated_ids": max(0, len(created_ids)-100)}






@app.post("/admin/reindex")
def admin_reindex(request: Request):
    require_admin(request)
    try:
        n = rebuild_index_from_csv()
        return {"ok": True, "count": n}
    except Exception as e:
        return JSONResponse(status_code=500, content={"ok": False, "error": str(e)})

# [FN:admin_reindex_docs] ----------------------------------------------------
@app.post("/admin/reindex_docs")
def admin_reindex_docs(request: Request):
    require_admin(request)
    result = rebuild_docs_index_from_folder()
    return result


# [FN:_fetch_contact_from_sqlite] ----------------------------------------------------
def _fetch_contact_from_sqlite(fac_query: str, role_query: str) -> tuple[str, str] | None:
    """
    Return (role_label, value_text) for a single contact role at a given facility.
    role_query is normalized (e.g., 'admin' -> 'Administrator').
    value_text is something like: 'Isla – email: x@y.com, phone: 555-1234' (formats what exists).
    """
    if not fac_query or not role_query:
        return None

    # 1) Find the facility_id from a fuzzy name
    conn = _db()
    cur = conn.cursor()
    cur.execute("SELECT facility_id, facility_name FROM facilities;")
    rows = cur.fetchall()

    # simple fuzzy pick (reuse your existing normalization)
    def _canon(s): 
        return re.sub(r"[^a-z0-9]+", "", (s or "").lower())

    cq = _canon(fac_query)
    pick = None
    for fid, fname in rows:
        if _canon(fname) == cq or _canon(fid) == cq:
            pick = (fid, fname); break
        if cq and cq in _canon(fname):
            pick = pick or (fid, fname)

    if not pick:
        return None

    fid, _fname = pick

    # 2) normalize the role (Administrator, Director of Nursing, Admissions, etc.)
    #    There is already a role cleaner/hint in this file we can lean on:
    role_norm_map = {
        "admin": "Administrator",
        "administrator": "Administrator",
        "don": "Director of Nursing",
        "director of nursing": "Director of Nursing",
        "admissions": "Admissions",
        "marketing": "Marketing",
        "social services": "Social Services",
        "social worker": "Social Services",
        "therapy": "Therapy",
        "physician": "Physician/NP",
        "doctor": "Physician/NP",
        "md": "Physician/NP",
        "np": "Physician/NP",
        "medical director": "Medical Director",
    }
    rq = (role_query or "").strip().lower()
    role_label = role_norm_map.get(rq, role_query.strip().title())

    # 3) fetch contact row for that role
    cur.execute("""
        SELECT type, name, email, phone, pref
        FROM facility_contacts
        WHERE facility_id = ?
        ORDER BY id ASC;
    """, (fid,))
    rows = cur.fetchall()

    # allow loose matching on 'type'
    def _match_type(t: str) -> bool:
        t = (t or "").strip().lower()
        return (t == rq) or (t == role_label.lower())

    for t, name, email, phone, pref in rows:
        if _match_type(t):
            bits = []
            if name:  bits.append(str(name).strip())
            # inline the contact fields that exist
            details = []
            if email: details.append(f"email: {email}")
            if phone: details.append(f"phone: {phone}")
            if pref:  details.append(f"pref: {pref}")
            if details:
                bits.append(" – " + ", ".join(details))
            return (role_label, "".join(bits))

    # no exact role match; try a very loose contains (e.g., t startswith)
    for t, name, email, phone, pref in rows:
        if (t or "").strip().lower().startswith(rq[:3]):
            bits = []
            if name:  bits.append(str(name).strip())
            details = []
            if email: details.append(f"email: {email}")
            if phone: details.append(f"phone: {phone}")
            if pref:  details.append(f"pref: {pref}")
            if details:
                bits.append(" – " + ", ".join(details))
            return (role_label, "".join(bits))

    return None


# ==============================================================================
# [S17] INTAKE SUBMIT — Generate QAs & upsert index
# ==============================================================================
# ==== Intake submit + Q&A generation =========================================
import os, csv, re
import pandas as pd
from fastapi import Body
try:
    from slugify import slugify  # pip install python-slugify
except Exception:
    # very small fallback if slugify isn't installed
    def slugify(x):
        return re.sub(r'[^a-z0-9]+', '-', str(x).strip().lower()).strip('-')

CLIENT_SHEET_PATH = os.path.join("data", "Client_Info_Spreadsheet.csv")
QA_CSV_PATH       = os.path.join("data", "qa.csv")

TOKEN_RE = re.compile(r"\[([^\]]+)\]")


# [FN:_best_facility_name] ----------------------------------------------------
def _best_facility_name(row_dict: dict) -> str:
    for k in ("legal_name", "facility_name", "doing_business_as", "dba"):
        if k in row_dict and str(row_dict[k]).strip():
            return str(row_dict[k]).strip()
    return ""


# [FN:_render_question] ----------------------------------------------------
def _render_question(template: str, row_dict: dict, header_map: dict) -> str:
    """Replace [Tokens] in Row-2 template using the submission row values."""
    def repl(m):
        token = m.group(1).strip()
        if token.lower() in ("facility", "facility name", "client"):
            val = _best_facility_name(row_dict)
            return val if val else m.group(0)
        key = header_map.get(token.lower())
        if key and str(row_dict.get(key, "")).strip():
            return str(row_dict[key]).strip()
        return m.group(0)
    return TOKEN_RE.sub(repl, str(template or ""))


# [FN:_make_qid] ----------------------------------------------------
def _make_qid(facility: str, field_name: str) -> str:
    return f"{slugify(facility) or 'client'}:{field_name}"


# [FN:generate_qas_from_sheet] ----------------------------------------------------
def generate_qas_from_sheet(sheet_csv_path: str, qa_csv_path: str):
    """
    Reads Client_Info_Spreadsheet.csv and upserts rows into qa.csv
    - Row 1: field names
    - Row 2: question templates
    - Row 3+: answers
    - Column B = Section (usually 'Facility Details')
    - Column C = Tags
    """
    df = pd.read_csv(sheet_csv_path, header=None)
    if df.shape[0] < 3:
        return

    row1 = df.iloc[0].fillna("").astype(str).tolist()  # field names
    row2 = df.iloc[1].fillna("").astype(str).tolist()  # question templates
    
    # NEW: Optional Row 3 = per-column tags
    row3_tags = df.iloc[2].fillna("").astype(str).tolist() if df.shape[0] >= 3 else []
    # Default: submissions start on Row 4 if Row 3 is a tags row (B/C empty)
    start_row = 4 if (row3_tags and row3_tags[1] == "" and row3_tags[2] == "") else 3

    # Build case-insensitive map e.g. [City] -> 'city'
    header_map = {}
    for name in row1:
        if not name: 
            continue
        header_map[name.strip().lower()] = name

    # Existing qa.csv (for upsert by id)
    try:
        qa = pd.read_csv(qa_csv_path).fillna("")
    except FileNotFoundError:
        qa = pd.DataFrame(columns=["id", "section", "question", "answer", "tags"])

    out_rows = []
    for _, r in df.iloc[start_row-1:].iterrows():  # submissions
        ...
        START_COL = 3  # D
        for col_idx, template in enumerate(row2[START_COL:], start=START_COL):
            if not template:
                continue
            if col_idx >= len(row1):
                continue
            field_name = row1[col_idx]
            ans = str(r.iloc[col_idx]) if col_idx < len(r) else ""
            if not str(ans).strip():
                continue

            q = _render_question(template, row_dict, header_map)
            qid = _make_qid(facility, field_name)

            # Keep Column C (global tags) + optional Row 3 per-column tag (if present)
            tags_this = tags
            if row3_tags and col_idx < len(row3_tags) and str(row3_tags[col_idx]).strip():
                t_extra = str(row3_tags[col_idx]).strip()
                tags_this = ", ".join([t for t in [tags, t_extra] if t]).strip(", ").strip()

            out_rows.append({
                "id": qid,
                "section": section or "Facility Details",
                "question": q,
                "answer": ans,
                "tags": tags_this
            })



    if out_rows:
        new_df = pd.DataFrame(out_rows)
        if not qa.empty:
            qa = qa[~qa["id"].isin(new_df["id"])]
        qa = pd.concat([new_df, qa], ignore_index=True)
        qa.to_csv(qa_csv_path, index=False)


# [FN:_build_tags_from_payload] ----------------------------------------------------
def _build_tags_from_payload(payload: dict) -> str:
    if "__tag_hint" in payload and str(payload["__tag_hint"]).strip():
        return str(payload["__tag_hint"]).strip()
    name = (payload.get("legal_name") or payload.get("facility_name") or "").strip()
    bits = []
    if name:
        bits.append(f"f-{name}")
    return ", ".join([b for b in bits if b])


from fastapi import Body

@app.post("/intake/submit")
async def intake_submit(payload: dict = Body(...)):
    """
    1) Append a new row to data/Client_Info_Spreadsheet.csv (Row1 names, Row2 templates must exist)
    2) Stamp Column B (section) = 'Facility Details' and Column C (tags)
    3) Generate ONLY this submission's Q&A rows
    4) Upsert ONLY those rows into qa.csv and the qa_main Chroma index
    """
    # Basic validation
    facility_name = (payload.get("facility_name") or payload.get("legal_name") or "").strip()
    if not facility_name:
        return {"ok": False, "error": "facility_name (or legal_name) is required"}

    # Ensure sheet exists with at least two header rows
    if not os.path.exists(CLIENT_SHEET_PATH):
        return {"ok": False, "error": f"Missing {CLIENT_SHEET_PATH}. Put your sheet (with Row1 names & Row2 questions) there."}

    df = pd.read_csv(CLIENT_SHEET_PATH, header=None)
    if df.shape[0] < 2:
        return {"ok": False, "error": "Client_Info_Spreadsheet.csv must have Row 1 field names and Row 2 question templates"}

    row1 = df.iloc[0].fillna("").astype(str).tolist()
    row2 = df.iloc[1].fillna("").astype(str).tolist()

    # Build a blank new row sized to your sheet
    new_row = [""] * len(row1)

    # Column B and C
    section_value = "Facility Details"
    tags_value    = _build_tags_from_payload(payload)
    if len(new_row) > 1:
        new_row[1] = section_value
    if len(new_row) > 2:
        new_row[2] = tags_value
        
    # === NEW: upsert the normalized facility + children into SQLite ===
    try:
        fac_id = upsert_facility_and_children(payload)
    except HTTPException as e:
        # bubble up with your existing error style
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DB upsert failed: {e}")

    

    # Map payload keys into the new row using Row1 field names
    name_to_idx = {row1[i]: i for i in range(len(row1)) if row1[i]}
    for k, v in payload.items():
        if k in name_to_idx:
            new_row[name_to_idx[k]] = str(v).strip()

    # Append to the sheet
    with open(CLIENT_SHEET_PATH, "a", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(new_row)

    # Build a pd.Series for just-submitted row (to reuse the same logic)
    sub_series = pd.Series(new_row)

    # Generate Q&A rows ONLY for this submission
    submission_qas = _generate_qas_for_submission(row1, row2, sub_series)

    if USE_SQLITE_WRITE:
        for r in submission_qas:
            r_sql = dict(r); r_sql["topics"] = r_sql.get("section","")
            upsert_qa_sqlite(r_sql)
        if MIRROR_SQLITE_WRITES:
            _upsert_into_qa_csv(submission_qas)
    else:
        _upsert_into_qa_csv(submission_qas)
        if MIRROR_SQLITE_WRITES:
            for r in submission_qas:
                r_sql = dict(r); r_sql["topics"] = r_sql.get("section","")
                upsert_qa_sqlite(r_sql)

    _upsert_into_chroma(submission_qas)

    return {
        "ok": True,
        "created": len(submission_qas),
        "section": section_value,
        "tags": tags_value,
        "ids": [r["id"] for r in submission_qas],
    }



# ==============================================================================
# [S18] MAIN ENTRYPOINT (__main__)
# ==============================================================================
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=False)


# ==============================================================================
# [S19] ABBREVIATION ADMIN ROUTES
# ==============================================================================
# ---- Abbreviation Admin routes ----
from fastapi import Query

@app.api_route("/admin/abv/list", methods=["GET", "OPTIONS"])

# [FN:admin_abv_list] ----------------------------------------------------
def admin_abv_list(request: Request, token: str = Query(default="")):
    # accept either header (require_admin) or token query param
    if request.method != "OPTIONS":
        try:
            require_admin(request)
        except HTTPException:
            _assert_admin(token)
    rows = _read_abv_rows()
    for r in rows:
        if "meaning" not in r and "expansion" in r:
            r["meaning"] = r.pop("expansion")
    return {"rows": rows}

@app.post("/admin/abv/add")

# [FN:admin_abv_add] ----------------------------------------------------
def admin_abv_add(
    request: Request,
    abbr: str = Form(...),
    meaning: str = Form(...),
    notes: str = Form(""),
    token: str = Form(default="")
):
    try:
        require_admin(request)
    except HTTPException:
        _assert_admin(token)

    # Normalize inputs
    abbr_s   = (abbr or "").strip()
    meaning_s= (meaning or "").strip()
    notes_s  = (notes or "").strip()
    if not abbr_s or not meaning_s:
        raise HTTPException(status_code=400, detail="abbr and meaning are required")

    # Duplicate check across both stores (CSV + SQLite) for safety
    csv_rows = _read_abv_rows()
    existing = { (r.get("abbr") or "").lower().strip() for r in csv_rows }
    try:
        for r in abv_rows_from_sqlite():
            existing.add((r.get("abbr") or "").lower().strip())
    except Exception:
        pass
    if abbr_s.lower() in existing:
        raise HTTPException(status_code=409, detail="Abbreviation already exists")

    row_obj = {"abbr": abbr_s, "meaning": meaning_s, "notes": notes_s}

    if USE_SQLITE_WRITE:
        # Primary write → SQLite
        upsert_abbrev_sqlite(row_obj)
        # Optional mirror → CSV
        if MIRROR_SQLITE_WRITES:
            csv_rows.insert(0, row_obj)
            _write_abv_rows(csv_rows)
    else:
        # Primary write → CSV (legacy)
        csv_rows.insert(0, row_obj)
        _write_abv_rows(csv_rows)
        # Optional mirror → SQLite
        if MIRROR_SQLITE_WRITES:
            try:
                upsert_abbrev_sqlite(row_obj)
            except Exception as e:
                print("mirror abv add->sqlite failed:", e)

    return {"ok": True}


@app.post("/admin/abv/update")

# [FN:admin_abv_update] ----------------------------------------------------
def admin_abv_update(
    request: Request,
    abbr: str = Form(...),         # key
    meaning: str = Form(...),
    notes: str = Form(""),
    token: str = Form(default="")
):
    try:
        require_admin(request)
    except HTTPException:
        _assert_admin(token)

    abbr_s    = (abbr or "").strip()
    meaning_s = (meaning or "").strip()
    notes_s   = (notes or "").strip()
    if not abbr_s:
        raise HTTPException(status_code=400, detail="abbr is required")

    # Load CSV to keep legacy file in sync when mirroring
    rows = _read_abv_rows()
    found = False
    for r in rows:
        if (r.get("abbr") or "").lower().strip() == abbr_s.lower():
            r["meaning"] = meaning_s
            r["notes"]   = notes_s
            found = True
            break
    if not found:
        # Even if CSV doesn't have it, allow SQLite primary to upsert (acts like add)
        pass

    row_obj = {"abbr": abbr_s, "meaning": meaning_s, "notes": notes_s}

    if USE_SQLITE_WRITE:
        # Primary write → SQLite
        upsert_abbrev_sqlite(row_obj)
        # Optional mirror → CSV
        if MIRROR_SQLITE_WRITES:
            if not found:
                rows.insert(0, row_obj)
            _write_abv_rows(rows)
    else:
        # Primary write → CSV
        if not found:
            rows.insert(0, row_obj)
        _write_abv_rows(rows)
        # Optional mirror → SQLite
        if MIRROR_SQLITE_WRITES:
            try:
                upsert_abbrev_sqlite(row_obj)
            except Exception as e:
                print("mirror abv update->sqlite failed:", e)

    return {"ok": True}


@app.post("/admin/abv/delete")

# [FN:admin_abv_delete] ----------------------------------------------------
def admin_abv_delete(
    request: Request,
    abbr: str = Form(...),
    token: str = Form(default="")
):
    try:
        require_admin(request)
    except HTTPException:
        _assert_admin(token)

    abbr_s = (abbr or "").strip()
    if not abbr_s:
        raise HTTPException(status_code=400, detail="abbr is required")

    rows = _read_abv_rows()
    new_rows = [r for r in rows if (r.get("abbr") or "").lower().strip() != abbr_s.lower()]
    not_found = (len(new_rows) == len(rows))

    if USE_SQLITE_WRITE:
        # Primary delete → SQLite
        _exec_write("DELETE FROM abbreviations WHERE abbr = ?", (abbr_s,))
        # Optional CSV mirror
        if MIRROR_SQLITE_WRITES:
            if not not_found:
                _write_abv_rows(new_rows)
            else:
                # If CSV didn't have it, still ensure file is present
                _write_abv_rows(rows)  # no-op write
    else:
        # Primary delete → CSV
        if not_found:
            raise HTTPException(status_code=404, detail="Abbreviation not found")
        _write_abv_rows(new_rows)
        # Optional mirror → SQLite
        if MIRROR_SQLITE_WRITES:
            try:
                _exec_write("DELETE FROM abbreviations WHERE abbr = ?", (abbr_s,))
            except Exception as e:
                print("mirror abv delete->sqlite failed:", e)

    return {"ok": True}

@app.get("/admin/dictionary/list")
def admin_dictionary_list(request: Request, q: str = Query(default="")):
    require_admin(request)
    rows = dict_rows_from_sqlite()
    ql = (q or "").strip().lower()
    if ql:
        rows = [r for r in rows if ql in r["key"].lower() or ql in r["canonical"].lower()]
    return {"ok": True, "items": rows}

@app.post("/admin/dictionary/upsert")
def admin_dictionary_upsert(
    request: Request,
    key: str = Form(...),
    canonical: str = Form(...),
    kind: str = Form("abbr"),
    notes: str = Form(""),
    match_mode: str = Form("exact"),
):
    require_admin(request)
    upsert_dictionary_sqlite({
        "key": key, "canonical": canonical, "kind": kind,
        "notes": notes, "match_mode": match_mode
    })
    return {"ok": True}

@app.post("/admin/dictionary/delete")
def admin_dictionary_delete(request: Request, key: str = Form(...)):
    require_admin(request)
    _exec_write("DELETE FROM dictionary WHERE key = ?", (key.strip().lower(),))
    clear_dictionary_caches()
    return {"ok": True}


